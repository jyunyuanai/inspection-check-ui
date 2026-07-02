from __future__ import annotations

import hashlib
import hmac
import os
import re
import shutil
import subprocess
import uuid
import zipfile
from copy import deepcopy
from datetime import datetime, date
from pathlib import Path
from typing import Any
from io import BytesIO

import pandas as pd
import streamlit as st
from PIL import Image, ImageOps
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 工程抽查 UI 工具 - 第九十九版
#
# 目標：
# 1. 使用者把照片直接拖進照片表格
# 2. 使用者自行選日期、填檢查項目、設計值、實測值
# 3. 不做 自動辨識、不做自動辨識、不做辨識快取
# 4. 一鍵產出工程抽查照片 Word，可接續既有照片 Word / 抽查表 Word
#
# 執行：
# uv run --with streamlit --with pandas --with pillow --with python-docx --with pywin32 streamlit run inspection_check.py
# ============================================================


BASE_DIR = Path(__file__).resolve().parent

PHOTO_TEMPLATE_DIR = BASE_DIR / "工程抽查照片樣板"
DEFAULT_RECORD_TEMPLATE_DIR = BASE_DIR / "抽查紀錄表樣板"
DEFAULT_OUTPUT_DIR = BASE_DIR / "輸出"
DEFAULT_UPLOAD_CACHE_DIR = BASE_DIR / "UI匯入照片暫存"
RUNTIME_ROOT_DIR = BASE_DIR / ".streamlit_runtime"


IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}

DOC_EXTS = {".doc", ".docx"}
MAX_IMAGE_UPLOAD_BYTES = 200 * 1024 * 1024
MAX_DOC_UPLOAD_BYTES = 50 * 1024 * 1024


def find_libreoffice_binary() -> str | None:
    return shutil.which("soffice") or shutil.which("libreoffice")


def supports_legacy_doc_conversion() -> bool:
    """
    .doc 轉 .docx 需要 Windows + Microsoft Word COM / pywin32。
    雲端部署通常是 Linux，會透過 LibreOffice 將 .doc 自動轉成 .docx。
    """
    return os.name == "nt" or find_libreoffice_binary() is not None


def allowed_record_doc_exts() -> set[str]:
    return DOC_EXTS


def allowed_record_upload_types() -> list[str]:
    return ["doc", "docx"]


def ensure_session_runtime_paths() -> None:
    """
    每個瀏覽器 session 使用自己的暫存、輸出、抽查表樣板資料夾。
    這是雲端多人共用時避免互相覆蓋/清檔的第一層保護。
    """
    if "runtime_session_id" not in st.session_state:
        st.session_state["runtime_session_id"] = uuid.uuid4().hex

    sid = str(st.session_state["runtime_session_id"])
    session_root = RUNTIME_ROOT_DIR / sid

    st.session_state["runtime_root_dir"] = str(session_root)
    st.session_state["record_template_dir"] = str(session_root / "抽查紀錄表樣板")
    st.session_state["output_dir"] = str(session_root / "輸出")
    st.session_state["upload_cache_dir"] = str(session_root / "UI匯入照片暫存")
    st.session_state["image_work_dir"] = str(session_root / "輸出" / "_ui_word_images")


def record_template_dir() -> Path:
    return Path(str(st.session_state.get("record_template_dir", DEFAULT_RECORD_TEMPLATE_DIR)))


def output_dir() -> Path:
    return Path(str(st.session_state.get("output_dir", DEFAULT_OUTPUT_DIR)))


def upload_cache_dir() -> Path:
    return Path(str(st.session_state.get("upload_cache_dir", DEFAULT_UPLOAD_CACHE_DIR)))


def image_work_dir() -> Path:
    return Path(str(st.session_state.get("image_work_dir", output_dir() / "_ui_word_images")))


def cleanup_unused_project_folders() -> None:
    """
    第九十八版：
    使用者確認照片暫存要放回專案資料夾內，方便查找。

    仍不再主動建立：
    - 原始照片
    - 正確範本

    這兩個資料夾只有在空資料夾時才刪除；若裡面有檔案，保留避免誤刪。
    """
    for name in ["原始照片", "正確範本"]:
        p = BASE_DIR / name
        if p.exists() and p.is_dir():
            try:
                if not any(p.iterdir()):
                    p.rmdir()
            except Exception:
                pass


def init_dirs() -> None:
    # 保留目前必要資料夾，照片暫存也放回專案資料夾。
    for folder in [PHOTO_TEMPLATE_DIR, record_template_dir(), output_dir(), upload_cache_dir(), image_work_dir()]:
        folder.mkdir(parents=True, exist_ok=True)

    cleanup_unused_project_folders()


def is_safe_managed_dir(folder: Path) -> bool:
    """
    只允許清理程式自己管理的三個資料夾：
    - UI匯入照片暫存
    - 輸出
    - 輸出/_ui_word_images
    """
    try:
        folder = folder.resolve()
        allowed = {
            upload_cache_dir().resolve(),
            output_dir().resolve(),
            image_work_dir().resolve(),
            record_template_dir().resolve(),
        }
        return folder in allowed
    except Exception:
        return False


def clear_directory_contents(folder: Path) -> None:
    """
    清空資料夾內容，但保留資料夾本身。
    """
    if not is_safe_managed_dir(folder):
        raise ValueError(f"不允許清理非程式管理資料夾：{folder}")

    folder.mkdir(parents=True, exist_ok=True)

    for child in list(folder.iterdir()):
        try:
            if child.is_dir():
                shutil.rmtree(child)
            else:
                child.unlink()
        except Exception:
            # 檔案如果正被 Word 或系統占用，略過不讓程式中斷。
            pass


def clear_output_runtime_files() -> None:
    """
    產出 Word 前清空輸出資料夾。
    目的：輸出資料夾只保留本次最新 Word，不再越堆越多。
    """
    clear_directory_contents(output_dir())
    image_work_dir().mkdir(parents=True, exist_ok=True)


def clear_upload_cache_files() -> None:
    """
    清空 UI 匯入照片暫存。
    """
    clear_directory_contents(upload_cache_dir())


def clear_current_case_files() -> None:
    """
    第九十八版：
    切換新案件時，把會累積的資料夾用取代方式清掉：
    - UI匯入照片暫存
    - 輸出
    - 輸出/_ui_word_images
    """
    clear_upload_cache_files()
    clear_output_runtime_files()


def reset_current_case_state() -> None:
    """
    清空目前畫面上的照片資料與下載狀態。
    不清工程名稱、施工地點、施工廠商，方便先輸入新案資料。
    """
    st.session_state.pop("df", None)
    st.session_state.pop("uploaded_key", None)
    st.session_state.pop("last_generated_words", None)
    st.session_state.pop("existing_photo_word_bytes", None)
    st.session_state.pop("existing_photo_word_name", None)
    st.session_state.pop("existing_record_word_bytes", None)
    st.session_state.pop("existing_record_word_name", None)

    for key in list(st.session_state.keys()):
        if str(key).startswith("slot_uploaded_key_"):
            st.session_state.pop(key, None)

    st.session_state["extra_empty_slots"] = 0


def prune_upload_cache_to_current_df() -> None:
    """
    刪掉 UI匯入照片暫存 裡已經不在目前照片表格使用的照片。
    單張刪除照片後，也不會讓暫存資料夾越來越大。
    """
    try:
        keep_paths: set[Path] = set()
        df = st.session_state.get("df")
        if isinstance(df, pd.DataFrame) and "照片路徑" in df.columns:
            for value in df["照片路徑"].dropna().astype(str):
                try:
                    p = Path(value).resolve()
                    if p.exists() and upload_cache_dir().resolve() in p.parents:
                        keep_paths.add(p)
                except Exception:
                    pass

        upload_cache_dir().mkdir(parents=True, exist_ok=True)
        for child in list(upload_cache_dir().iterdir()):
            try:
                if child.is_file() and child.resolve() not in keep_paths:
                    child.unlink()
            except Exception:
                pass
    except Exception:
        pass


def file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def normalize_text(text: str) -> str:
    if not text:
        return ""

    replacements = {
        "：": ":",
        "，": ",",
        "。": ".",
        "　": " ",
        "／": "/",
        "－": "-",
        "—": "-",
        "～": "~",
        "㎝": "cm",
        "ＣＭ": "cm",
        "ｃｍ": "cm",
        "公分": "cm",
        "厘米": "cm",
        "０": "0",
        "１": "1",
        "２": "2",
        "３": "3",
        "４": "4",
        "５": "5",
        "６": "6",
        "７": "7",
        "８": "8",
        "９": "9",
        "Ｏ": "0",
        "ｏ": "0",
    }

    for k, v in replacements.items():
        text = text.replace(k, v)

    text = re.sub(r"A\s*[O0]\s*K", "A0K", text, flags=re.IGNORECASE)
    text = re.sub(r"A0K\s*\+\s*", "A0K+", text, flags=re.IGNORECASE)
    text = re.sub(r"A0K\+0?(\d{1,3})", lambda m: f"A0K+{int(m.group(1)):03d}", text)

    text = re.sub(
        r"\bD\s*(\d{2})\s*[O0@◎]\s*(\d{1,3})\s*(CM|cm)?",
        r"D\1@\2cm",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"\bD\s*(\d{2})\s*@\s*(\d{1,3})\s*(CM|cm)?",
        r"D\1@\2cm",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"(\d{1,4})\s*[cC][.\s]*[mM]", r"\1cm", text)
    text = text.replace("設 計", "設計").replace("實 測", "實測").replace("檢 查", "檢查")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def natural_sort_text(value: str) -> str:
    """
    第九十八版：
    補回自然排序函式，修正產出 Word 時：
    NameError: name 'natural_sort_text' is not defined

    用途：
    讓照片檔名 / 分組 / 檢查項目排序時，數字可以照 1、2、10 的順序排，
    不會變成 1、10、2。
    """
    text = normalize_text(str(value or ""))

    parts = re.split(r"(\d+)", text)
    out_parts: list[str] = []

    for part in parts:
        if part.isdigit():
            out_parts.append(part.zfill(10))
        else:
            out_parts.append(part.lower())

    return "".join(out_parts)


def parse_date_from_text(text: str) -> str:
    """
    第九十八版：
    修正照片右下角 2026/5/21 被誤判成 126年5月21日的問題。

    支援：
    - 2026/5/21、2026-05-21、2026.5.21  → 115年5月21日
    - 20260521                              → 115年5月21日
    - 115年5月21日、115/5/21、115-05-21     → 115年5月21日
    - 1150521                               → 115年5月21日
    """
    t = normalize_text(text)

    # 1) 西元 4 碼：2026/5/21、2026-05-21、2026.5.21
    m = re.search(r"(?<!\d)(\d{4})[./\-](\d{1,2})[./\-](\d{1,2})(?!\d)", t)
    if m:
        year = int(m.group(1))
        month = int(m.group(2))
        day = int(m.group(3))
        if 1911 <= year <= 2100 and 1 <= month <= 12 and 1 <= day <= 31:
            return f"{year - 1911}年{month}月{day}日"

    # 2) 西元 8 碼：20260521
    m = re.search(r"(?<!\d)(20\d{2})(\d{2})(\d{2})(?!\d)", t)
    if m:
        year = int(m.group(1))
        month = int(m.group(2))
        day = int(m.group(3))
        if 1911 <= year <= 2100 and 1 <= month <= 12 and 1 <= day <= 31:
            return f"{year - 1911}年{month}月{day}日"

    # 3) 民國：115年5月21日
    m = re.search(r"(?<!\d)(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", t)
    if m:
        roc_year = int(m.group(1))
        month = int(m.group(2))
        day = int(m.group(3))
        if roc_year < 100:
            roc_year += 100
        if 1 <= month <= 12 and 1 <= day <= 31:
            return f"{roc_year}年{month}月{day}日"

    # 4) 民國斜線：115/5/21、115-05-21
    #    用 (?<!\d) 避免從 2026/5/21 裡面吃到 026/5/21。
    m = re.search(r"(?<!\d)(\d{2,3})[./\-](\d{1,2})[./\-](\d{1,2})(?!\d)", t)
    if m:
        roc_year = int(m.group(1))
        month = int(m.group(2))
        day = int(m.group(3))
        if roc_year < 100:
            roc_year += 100
        if 1 <= month <= 12 and 1 <= day <= 31:
            return f"{roc_year}年{month}月{day}日"

    # 5) 民國 7 碼：1150521
    m = re.search(r"(?<!\d)(\d{7})(?!\d)", t)
    if m:
        raw = m.group(1)
        roc_year = int(raw[:3])
        month = int(raw[3:5])
        day = int(raw[5:7])
        if roc_year < 100:
            roc_year += 100
        if 1 <= month <= 12 and 1 <= day <= 31:
            return f"{roc_year}年{month}月{day}日"

    return ""


def roc_date_sort_value(value: str) -> int:
    """
    第九十八版：
    Word 輸出時依日期排序。
    支援民國與西元日期。
    """
    text = normalize_text(str(value or ""))

    m = re.search(r"(?<!\d)(\d{4})[./\-](\d{1,2})[./\-](\d{1,2})(?!\d)", text)
    if m:
        try:
            y = int(m.group(1))
            mo = int(m.group(2))
            d = int(m.group(3))
            if 1911 <= y <= 2100 and 1 <= mo <= 12 and 1 <= d <= 31:
                return y * 10000 + mo * 100 + d
        except Exception:
            pass

    m = re.search(r"(?<!\d)(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", text)
    if not m:
        m = re.search(r"(?<!\d)(\d{2,3})[./\-](\d{1,2})[./\-](\d{1,2})(?!\d)", text)

    if not m:
        return 99999999

    try:
        y = int(m.group(1))
        mo = int(m.group(2))
        d = int(m.group(3))
        if y < 100:
            y += 100
        if not (1 <= mo <= 12 and 1 <= d <= 31):
            return 99999999

        current_roc = date.today().year - 1911
        if y > current_roc + 8 and (y - 11) >= 100:
            y -= 11

        return (y + 1911) * 10000 + mo * 100 + d
    except Exception:
        return 99999999


def clean_photo_work_item_label(value: Any) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    invalid_parts = ["請選擇", "請先匯入", "待確認"]
    if any(x in text for x in invalid_parts):
        return ""

    # 移除附錄頁碼與來源描述，保留主要工項。
    text = re.sub(r"（.*?）", "", text)
    text = re.sub(r"\|.*$", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def photo_word_work_item_label(row: pd.Series) -> str:
    """
    第九十八版：
    工程抽查照片 Word 分頁用工項。

    修正：
    之前若沒有選「抽查表工項」，程式會拿「檢查項目」當分頁工項。
    這會造成同一天、同一類工程的兩張照片，只因檢查項目文字不同就被拆成兩頁。

    現在改成：
    1. 有選抽查表工項 / 抽查表名稱 / 對應表單 → 依該工項分頁。
    2. 沒選工項 → 統一視為「未分類」，同日期最多兩張可放同一頁。
    """
    for key in ["抽查表工項", "抽查表名稱", "對應表單"]:
        label = clean_photo_work_item_label(row.get(key, ""))
        if label:
            return label

    return "未分類"


def photo_word_page_group_key(row: pd.Series) -> tuple:
    """
    第九十八版：
    工程抽查照片 Word 頁面分組規則：
    - 同日期 + 同工項：可以放同一頁，最多兩張
    - 同日期 + 不同工項：分不同頁
    - 不同日期：分不同頁
    - 沒選工項：同日期統一當同一組，不會因檢查項目不同而拆頁
    """
    date_text = str(row.get("日期", "") or "").strip()
    item_text = photo_word_work_item_label(row)
    return (roc_date_sort_value(date_text), date_text, natural_sort_text(item_text))


def sort_for_word_output(df: pd.DataFrame) -> pd.DataFrame:
    """
    Word 輸出排序：
    1. 日期由早到晚
    2. 同一天依照片 Word 工項分組
       未選工項者統一視為「未分類」，不再用檢查項目拆頁
    3. 同一天同工項依檢查項目 / 樁號排序
    4. 最後依照片檔名
    """
    out = df.copy()
    out["_日期排序"] = out["日期"].apply(roc_date_sort_value)
    out["_照片頁面工項排序"] = out.apply(lambda r: natural_sort_text(photo_word_work_item_label(r)), axis=1)

    if "照片分組" in out.columns:
        out["_分組排序"] = out["照片分組"].astype(str).apply(natural_sort_text)
    else:
        out["_分組排序"] = ""

    out["_項目排序"] = out["檢查項目"].astype(str).apply(natural_sort_text)
    out["_照片排序"] = out["照片檔名"].astype(str).apply(natural_sort_text)

    out = out.sort_values(
        ["_日期排序", "_照片頁面工項排序", "_分組排序", "_項目排序", "_照片排序"],
        ascending=[True, True, True, True, True],
        kind="mergesort",
    )

    return out.drop(columns=["_日期排序", "_照片頁面工項排序", "_分組排序", "_項目排序", "_照片排序"], errors="ignore")


def extract_group_name(path: Path) -> str:
    """
    第九十八版：
    已取消掃描「原始照片」資料夾。
    這裡只保留安全的分組名稱判斷；UI 暫存照片不另外分組。
    """
    parent = Path(path).parent.name
    if parent in ["UI匯入照片暫存", "工程抽查UI匯入照片暫存"]:
        return ""
    return parent


def table_text(table) -> str:
    """
    把 Word 表格文字攤平成字串，用來判斷這張表是不是 B05 / B06 等工項。
    """
    parts: list[str] = []
    try:
        for r in table.rows:
            for c in r.cells:
                text = normalize_text(c.text)
                if text:
                    parts.append(text)
    except Exception:
        pass
    return "\n".join(parts)


def is_catalog_table_text(text: str) -> bool:
    """
    判斷是不是「工程抽查紀錄表目錄」那張目錄表。
    目錄表只拿來產生選單，不可以被當作抽查紀錄表輸出。
    """
    t = normalize_text(text)
    b_count = len(re.findall(r"\bB\d{2}\b", t, re.IGNORECASE))
    if "抽查程序流程圖名稱" in t and "頁碼" in t:
        return True
    if "附錄二" in t and b_count >= 3:
        return True
    if "工程抽查紀錄表目錄" in t:
        return True
    return False


def score_record_table_for_row(table, row: pd.Series) -> int:
    """
    分數越高，越像照片卡片選到的那一張抽查表。
    例如卡片選 B05 灌溉溝，就優先找表格內有「編號：B05」與「灌溉溝」的表。
    """
    text = table_text(table)
    upper = text.upper()
    code = str(row.get("抽查表編號", "") or record_work_item_code_from_label(str(row.get("抽查表工項", ""))) or "").upper()
    name = str(row.get("抽查表名稱", "") or record_work_item_name_from_label(str(row.get("抽查表工項", ""))) or "")
    label = str(row.get("抽查表工項", "") or "")

    if not text:
        return -9999

    score = 0

    # 目錄表不能輸出。
    if is_catalog_table_text(text):
        score -= 1000

    if code:
        # 最重要：表格裡明確寫「編號：B05」。
        if re.search(rf"編\s*號\s*[：:]?\s*{re.escape(code)}", upper, re.IGNORECASE):
            score += 800
        if code in upper:
            score += 220

    name_key = re.sub(r"\s+", "", name)
    text_key = re.sub(r"\s+", "", text)
    if name_key and name_key in text_key:
        score += 450

    if "抽查紀錄表" in text:
        score += 50

    # 若工項名稱 / 下拉選單文字出現在表格裡，也補分。
    # 不再使用固定 B01~B10 清單，避免不同案件的工項被寫死。
    candidate_tokens: list[str] = []
    for source_text in [name, label]:
        source_text = re.sub(r"B\d{2}", "", str(source_text or ""), flags=re.IGNORECASE)
        source_text = re.sub(r"抽查紀錄表|工程|施工|附錄.*", "", source_text)
        for token in re.split(r"[\s_｜|/\\\-]+", source_text):
            token = token.strip()
            if len(token) >= 2 and token not in candidate_tokens:
                candidate_tokens.append(token)

    for token in candidate_tokens:
        if token and token in text:
            score += 100

    return score


def find_selected_record_table(template_path: Path, row: pd.Series):
    """
    從整份抽查紀錄表 Word 裡，只找出照片卡片所選工項的那一張表。
    回傳的是來源文件中的 table 物件；找不到時回傳 None。
    """
    from docx import Document

    try:
        src_doc = Document(str(template_path))
    except Exception:
        return None

    return find_selected_record_table_in_doc(src_doc, row)


def find_selected_record_table_in_doc(doc, row: pd.Series):
    best_table = None
    best_score = -9999

    for table in doc.tables:
        score = score_record_table_for_row(table, row)
        if score > best_score:
            best_score = score
            best_table = table

    # 分數太低代表沒有找到真正對應的抽查表。
    if best_score < 100:
        return None

    return best_table


def keep_only_table_in_document(doc, table) -> None:
    """
    Preserve the selected table inside its original document instead of copying it
    across documents. This keeps Word-only layout details such as vertical cells.
    """
    try:
        body = doc._body._element
        target_tbl = table._tbl
        sect = source_table_section_pr(table)
        fallback_sect = body.find(qn("w:sectPr"))
        if sect is None and fallback_sect is not None:
            sect = deepcopy(fallback_sect)

        for child in list(body):
            if child is target_tbl:
                continue
            body.remove(child)

        if sect is not None:
            body.append(deepcopy(sect))
    except Exception:
        pass


def clear_document_body_keep_section(doc) -> None:
    """
    第九十八版：
    抽查表 Word 要盡量跟原始樣本一樣，不能跑版。

    做法：
    直接用使用者匯入的原始 Word 當底稿，先清掉正文，
    但保留：
    - styles.xml 樣式
    - section margins / page size / header/footer 設定
    - 文件預設字型與表格樣式
    再把選中的 B05/B06 單張表格複製進來。
    """
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def remove_empty_body_paragraphs(doc) -> None:
    """
    第九十八版：
    工程抽查照片 Word 表格上方不保留多餘空白段落。
    只移除完全空白且沒有圖片/物件的段落，不刪除有文字的抬頭。
    """
    for p in list(doc.paragraphs):
        text = (p.text or "").strip()
        xml = p._element.xml

        has_drawing = "<w:drawing" in xml or "<w:pict" in xml
        has_page_break = 'w:type="page"' in xml or "w:br" in xml
        has_section = "<w:sectPr" in xml

        if not text and not has_drawing and not has_page_break and not has_section:
            try:
                p._element.getparent().remove(p._element)
            except Exception:
                pass


def compact_photo_word_spacing(doc) -> None:
    """
    讓工程抽查照片 Word 抬頭與表格不要空太大。
    第八十七版修正：補上 Pt 匯入，避免這段因 NameError 被 try/except 吃掉。
    """
    from docx.shared import Pt

    try:
        for p in doc.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
    except Exception:
        pass


def force_photo_title_size_16(doc) -> None:
    """
    第九十八版：
    工程抽查照片標題字大小一律 16pt。
    包含：
    - 程式新增的每頁標題
    - 樣板原本留下來的標題
    """
    from docx.shared import Pt

    try:
        for p in doc.paragraphs:
            if (p.text or "").strip() == "工程抽查照片":
                if not p.runs:
                    r = p.add_run("工程抽查照片")
                    set_run_font_kai(r)
                for r in p.runs:
                    r.font.size = Pt(16)
                    r.bold = True
                    set_run_font_kai(r)
    except Exception:
        pass


def append_copied_table(doc, source_table):
    """
    把來源 Word 的單一表格完整複製到新 Word。
    """
    from copy import deepcopy
    from docx.oxml.ns import qn
    from docx.table import Table

    new_tbl = deepcopy(source_table._tbl)
    body = doc._body._element
    sect_pr = body.find(qn("w:sectPr"))

    if sect_pr is not None:
        body.insert(body.index(sect_pr), new_tbl)
    else:
        body.append(new_tbl)

    return Table(new_tbl, doc)



def source_table_section_pr(source_table):
    """
    取出來源抽查表所在頁面的 section 設定。
    Word 的頁邊界/紙張設定常放在表格後面的段落 sectPr，
    只複製表格但沒複製 sectPr 時，表格可能會跑版或被擠到下一頁。
    """
    try:
        body = source_table._tbl.getparent()
        children = list(body)
        idx = children.index(source_table._tbl)

        # 先找表格後面最近的 sectPr，這通常就是該頁抽查表的版面設定。
        for el in children[idx + 1:]:
            sect = el.find('.//' + qn('w:sectPr'))
            if sect is not None:
                return deepcopy(sect)

        # 找不到時，用文件最後的 sectPr。
        sect = body.find(qn('w:sectPr'))
        if sect is not None:
            return deepcopy(sect)
    except Exception:
        pass
    return None


def apply_source_table_section_to_doc(doc, source_table) -> None:
    """
    讓輸出的抽查紀錄表使用來源表格原本的頁面設定。
    不改字、不改表格，只同步 Word section 設定。
    """
    try:
        src_sect = source_table_section_pr(source_table)
        if src_sect is None:
            return

        body = doc._body._element
        old_sect = body.find(qn('w:sectPr'))
        if old_sect is not None:
            body.remove(old_sect)
        body.append(deepcopy(src_sect))
    except Exception:
        pass

def remove_page_breaks_from_xml(root) -> None:
    """
    移除抽查表樣本裡可能藏著的手動分頁符號。
    這會避免簽名欄被硬切到下一頁。
    """
    try:
        for br in list(root.iter(qn("w:br"))):
            if br.get(qn("w:type")) == "page":
                parent = br.getparent()
                if parent is not None:
                    parent.remove(br)
        for tag in ["w:lastRenderedPageBreak"]:
            for el in list(root.iter(qn(tag))):
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
    except Exception:
        pass


def set_record_page_small_margins(doc) -> None:
    """
    抽查紀錄表輸出專用：把 A4 頁邊界縮小，增加可用高度與寬度。
    Word 會把表格後方的隱形段落也算進版面；樣板原檔可一頁，
    單獨抽出表格後仍可能因頁邊界不足而把簽名列推到第二頁。
    """
    try:
        from docx.shared import Cm
        for section in doc.sections:
            section.top_margin = Cm(0.7)
            section.bottom_margin = Cm(0.6)
            section.left_margin = Cm(0.8)
            section.right_margin = Cm(0.8)
            section.header_distance = Cm(0.3)
            section.footer_distance = Cm(0.3)
    except Exception:
        pass


def force_record_table_fit_one_page(doc, table) -> None:
    """
    第九十八版：保守修正。

    前一版把抽查表字級與列高硬壓縮，
    會把使用者原本的抽查表格式改壞。

    這版只做不影響版面的安全處理：
    - 移除樣本內可能殘留的手動分頁符號
    - 清除段落的「段前分頁 / 與下段同頁 / 段落中不分頁」設定
    - 不改字級
    - 不改表格寬度
    - 不改列高
    - 不改頁邊界
    """
    try:
        remove_page_breaks_from_xml(table._tbl)

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    try:
                        p.paragraph_format.keep_with_next = False
                        p.paragraph_format.keep_together = False
                        p.paragraph_format.page_break_before = False
                    except Exception:
                        pass

                for nested_table in cell.tables:
                    force_record_table_fit_one_page(doc, nested_table)
    except Exception:
        pass


def restore_record_stage_cells_layout(table) -> None:
    """
    Force the left-side stage labels to the requested stacked style:
    施 / 工 / 前, 施 / 工 / 中, 施 / 工 / 後.

    This avoids depending on Word textDirection, which can render differently
    after LibreOffice conversion or cross-document table copies.
    """
    stage_labels = {"施工前", "施工中", "施工後"}
    seen_cells = []

    try:
        from docx.shared import Cm, Pt

        for row in table.rows:
            for cell in row.cells:
                if any(cell._tc is seen for seen in seen_cells):
                    continue
                seen_cells.append(cell._tc)

                text = re.sub(r"\s+", "", normalize_text(cell.text))
                if text not in stage_labels:
                    continue

                tc_pr = cell._tc.get_or_add_tcPr()
                for tag in ["w:textDirection", "w:noWrap", "w:tcFitText"]:
                    old = tc_pr.find(qn(tag))
                    if old is not None:
                        tc_pr.remove(old)
                for old in list(cell._tc.iter(qn("w:textDirection"))):
                    parent = old.getparent()
                    if parent is not None:
                        parent.remove(old)

                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), "360")
                try:
                    cell.width = Cm(0.64)
                except Exception:
                    pass

                v_align = tc_pr.find(qn("w:vAlign"))
                if v_align is None:
                    v_align = OxmlElement("w:vAlign")
                    tc_pr.append(v_align)
                v_align.set(qn("w:val"), "center")

                clear_cell(cell)
                p = cell.paragraphs[0]
                try:
                    p.alignment = 1
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = Pt(10)
                except Exception:
                    pass
                for idx, char in enumerate(text):
                    run = p.add_run(char)
                    set_run_font_kai(run)
                    try:
                        run.font.size = Pt(10)
                    except Exception:
                        pass
                    if idx < len(text) - 1:
                        run.add_break()
    except Exception:
        pass


def optimize_record_table_output_layout(doc, table) -> None:
    """
    抽查紀錄表輸出前的保守版面整理。

    只處理 Word 容易造成跑版的版面上下文：
    - 套用抽查表輸出專用緊湊頁邊界
    - 移除手動分頁符號
    - 清掉段落 keep/page-break 設定

    不縮字、不硬改列高、不改表格文字內容。
    """
    try:
        set_record_page_small_margins(doc)
        force_record_table_fit_one_page(doc, table)
        restore_record_stage_cells_layout(table)
    except Exception:
        pass


def _set_tiny_record_tail_paragraph(paragraph_el) -> None:
    """
    把抽查表後方的空白結尾段落壓到 1pt。
    Word 文件以表格結尾時，仍會需要一個段落承接 section；
    這個段落若使用預設 12pt，容易形成空白尾頁。
    """
    try:
        p_pr = paragraph_el.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            paragraph_el.insert(0, p_pr)

        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        spacing.set(qn("w:line"), "1")
        spacing.set(qn("w:lineRule"), "exact")

        r_pr = p_pr.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            p_pr.append(r_pr)

        for tag in ["w:sz", "w:szCs"]:
            node = r_pr.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                r_pr.append(node)
            node.set(qn("w:val"), "2")
    except Exception:
        pass


def _is_blank_layout_paragraph(paragraph_el) -> bool:
    try:
        for text_el in paragraph_el.iter(qn("w:t")):
            if (text_el.text or "").strip():
                return False
        xml = paragraph_el.xml
        if "<w:drawing" in xml or "<w:pict" in xml or "<w:br" in xml:
            return False
        return True
    except Exception:
        return False


def ensure_compact_record_document_tail(doc) -> None:
    """
    避免抽查表輸出檔最後多出空白頁。

    python-docx 可以讓 body 直接以 table + sectPr 結尾，
    但 Word 實際開啟時仍會用預設段落高度計算結尾段落。
    在表格與 sectPr 之間補一個 1pt 空段落，可保留 Word 需要的結尾，
    又不會把空白段落擠到下一頁。
    """
    try:
        body = doc._body._element
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is None:
            return

        sect_idx = body.index(sect_pr)
        if sect_idx <= 0:
            tail_p = OxmlElement("w:p")
            body.insert(sect_idx, tail_p)
            _set_tiny_record_tail_paragraph(tail_p)
            return

        prev = body[sect_idx - 1]
        if prev.tag == qn("w:tbl"):
            tail_p = OxmlElement("w:p")
            body.insert(sect_idx, tail_p)
            _set_tiny_record_tail_paragraph(tail_p)
        elif prev.tag == qn("w:p") and _is_blank_layout_paragraph(prev):
            _set_tiny_record_tail_paragraph(prev)
    except Exception:
        pass

def record_output_unique_key(row: pd.Series) -> tuple[str, str]:
    """
    抽查紀錄表去重規則：同一天 + 同一個抽查表工項，只產生一張。
    優先用 B05/B07 編號；沒有編號時才用選單文字。
    """
    date_text = str(row.get("日期", "") or "").strip()
    code = str(row.get("抽查表編號", "") or record_work_item_code_from_label(str(row.get("抽查表工項", ""))) or "").strip().upper()
    if not code:
        code = clean_photo_work_item_label(str(row.get("抽查表工項", "") or row.get("抽查表名稱", "") or ""))
    return (str(roc_date_sort_value(date_text)), code or "未分類")


def append_page_break_before_next_table(doc) -> None:
    try:
        doc.add_page_break()
    except Exception:
        pass


def append_value_to_cell(cell, value: str) -> None:
    value = str(value or "").strip()
    if not value:
        return

    current = normalize_text(cell.text)
    if value in current:
        return

    try:
        if current:
            cell.add_paragraph(value)
        else:
            cell.text = value
    except Exception:
        try:
            cell.text = value
        except Exception:
            pass


def compact_single_line_value(value: str) -> str:
    """
    抽查表欄位取代用：
    把換行、Tab、多空白整理成單行，避免表格列高被撐高。
    """
    value = str(value or "").strip()
    value = re.sub(r"[\r\n\t]+", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value


def replace_value_to_cell(cell, value: str, force_font: str | None = None) -> None:
    """
    第九十八版：保留樣板原格式，只取代文字。

    重點：
    - 不改字級
    - 不改字型（除非呼叫端明確指定 force_font）
    - 不改段落行距
    - 不改表格寬度 / 列高 / 頁邊界
    - 盡量沿用原本儲存格第一個 run 的格式
    """
    value = compact_single_line_value(value)
    if not value:
        return

    try:
        if not cell.paragraphs:
            cell.add_paragraph("")

        p = cell.paragraphs[0]

        # 優先沿用原本第一個 run，這樣 Word 樣板的字級、字型、粗細不會被改掉。
        if p.runs:
            p.runs[0].text = value
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(value)

        # 其他段落若有舊值，清文字但不刪段落，避免破壞樣板表格結構。
        for extra_p in cell.paragraphs[1:]:
            for r in extra_p.runs:
                r.text = ""

        if force_font:
            for run in p.runs:
                set_run_font_kai(run, force_font)
    except Exception:
        try:
            cell.text = value
        except Exception:
            pass

def replace_value_in_same_label_cell(cell, label: str, value: str, force_font: str | None = None) -> None:
    """
    有些 Word 表格可能是同一格寫：工程名稱：舊名稱。
    沒有右邊空格時，保留標籤，把標籤後面的內容取代成新值。
    """
    value = compact_single_line_value(value)
    if not value:
        return

    label_text = str(label or "").strip()
    text = normalize_text(cell.text)

    if label_text and label_text in text:
        new_text = f"{label_text}：{value}"
    else:
        new_text = value

    replace_value_to_cell(cell, new_text, force_font=force_font)


def replace_neighbor_cell_by_label(table, labels: list[str], value: str, force_font: str | None = None) -> bool:
    """
    找到含有標籤的格子，優先把右邊空格「整格取代」成 value。
    若沒有右邊空格，才取代同一格標籤後面的內容。

    用於：
    - 工程名稱
    - 施工廠商
    """
    value = str(value or "").strip()
    if not value:
        return False

    for row in table.rows:
        cells = list(row.cells)
        seen_tc = set()

        for idx, cell in enumerate(cells):
            if id(cell._tc) in seen_tc:
                continue
            seen_tc.add(id(cell._tc))

            text = normalize_text(cell.text)
            matched_label = ""
            for label in labels:
                if label and label in text:
                    matched_label = label
                    break

            if not matched_label:
                continue

            # 優先填右邊不同 XML cell，避免 merged cell 重複。
            for j in range(idx + 1, len(cells)):
                if cells[j]._tc is not cell._tc:
                    replace_value_to_cell(cells[j], value, force_font=force_font)
                    return True

            # 若沒有右邊空格，只好同格取代。
            replace_value_in_same_label_cell(cell, matched_label, value, force_font=force_font)
            return True

    return False


def fill_selected_record_table(table, row: pd.Series) -> None:
    """
    第九十八版：
    抽查紀錄表仍然不填：
    日期、檢查項目、設計值、實測值。

    這次依使用者需求，以下欄位改成「取代」：
    1. 工程名稱：使用 UI 左側 / 表格內的工程名稱取代樣本舊值
    2. 施工廠商：使用 UI 左側 / 表格內的施工廠商取代樣本舊值

    不再用新增段落，所以不會變成「舊值 + 新值」。
    """
    project = str(row.get("工程名稱", "") or "").strip()
    contractor = str(row.get("施工廠商", "") or "").strip()

    if project:
        replace_neighbor_cell_by_label(table, ["工程名稱", "工程名"], project)

    if contractor:
        # 優先找施工廠商欄位；部分表格可能寫成承攬廠商或廠商。
        # 注意：「廠商」只用於施工廠商，不碰工程名稱。
        # 施工廠商一律套用標楷體，不沿用樣板原字型。
        replace_neighbor_cell_by_label(table, ["施工廠商", "承攬廠商", "廠商"], contractor, force_font="標楷體")

    # 重要：不要在這裡壓縮字級、行距、列高或表格寬度。
    # 抽查紀錄表必須照使用者提供的 Word 樣板原格式輸出。
    return


def create_record_words(
    df: pd.DataFrame,
    output_file_name: str = '',
    base_record_word_bytes: bytes | None = None,
) -> list[dict[str, Any]]:
    """
    第九十八版：
    - 支援接續舊抽查紀錄表 Word。
    - 同一天 + 同一個抽查表工項，只輸出一張抽查表。
    - 抽查紀錄表保留樣板字級、行距、列高與表格寬度。
    - 輸出時套用抽查表專用緊湊頁邊界，避免簽名列被 Word 推到第二頁。
    """
    from docx import Document

    selected = df[df["輸出"] == True].copy()

    if selected.empty:
        return []

    selected = sort_for_word_output(selected)
    rows_to_output = []
    seen_unique_keys: set[tuple[str, str]] = set()

    for _, row in selected.iterrows():
        label = str(row.get("抽查表工項", "") or "").strip()
        template_path = str(row.get("抽查表工項路徑", "") or row.get("抽查表樣本路徑", "") or "").strip()

        if not label or not template_path or not Path(template_path).exists():
            continue

        unique_key = record_output_unique_key(row)
        if unique_key in seen_unique_keys:
            continue
        seen_unique_keys.add(unique_key)

        source_table = find_selected_record_table(Path(template_path), row)
        if source_table is None:
            try:
                st.warning(f"找不到 {label} 對應的單張抽查表，已略過該照片。")
            except Exception:
                pass
            continue

        rows_to_output.append((row, source_table))

    if not rows_to_output:
        return []

    append_to_existing = bool(base_record_word_bytes)
    preserved_first_table = False

    if append_to_existing:
        out_doc = Document(BytesIO(base_record_word_bytes))
    else:
        first_template = Path(str(rows_to_output[0][0].get("抽查表工項路徑", "") or rows_to_output[0][0].get("抽查表樣本路徑", "")))
        out_doc = Document(str(first_template))
        first_row = rows_to_output[0][0]
        first_table = find_selected_record_table_in_doc(out_doc, first_row)

        if first_table is not None:
            keep_only_table_in_document(out_doc, first_table)
            fill_selected_record_table(first_table, first_row)
            optimize_record_table_output_layout(out_doc, first_table)
            preserved_first_table = True
        else:
            clear_document_body_keep_section(out_doc)

    rows_to_append = rows_to_output[1:] if preserved_first_table else rows_to_output

    for idx, (row, source_table) in enumerate(rows_to_append):
        if append_to_existing or preserved_first_table or idx > 0:
            append_page_break_before_next_table(out_doc)

        apply_source_table_section_to_doc(out_doc, source_table)
        new_table = append_copied_table(out_doc, source_table)
        fill_selected_record_table(new_table, row)
        optimize_record_table_output_layout(out_doc, new_table)

    selected_codes = []
    for row, _table in rows_to_output:
        code = str(row.get("抽查表編號", "") or record_work_item_code_from_label(str(row.get("抽查表工項", ""))) or "").strip()
        if code and code not in selected_codes:
            selected_codes.append(code)

    file_name = make_labeled_download_name(output_file_name, "抽查紀錄表")

    ensure_compact_record_document_tail(out_doc)

    # 抽查紀錄表不套用全文件字型，也不刪樣板空白段落；只做上方版面整理。
    return [{
        "file_name": file_name,
        "data": save_docx_to_bytes(out_doc),
    }]

def safe_open_image(path: Path) -> Image.Image:
    img = Image.open(path)
    img = ImageOps.exif_transpose(img)
    return img


def get_uploaded_file_size(uploaded_file: Any) -> int:
    size = getattr(uploaded_file, "size", None)
    if isinstance(size, int) and size >= 0:
        return size
    try:
        return len(uploaded_file.getbuffer())
    except Exception:
        return 0


def short_hash_bytes(data: bytes, length: int = 12) -> str:
    return hashlib.sha256(data).hexdigest()[:length]


def safe_upload_stem(filename: str, default_name: str = "upload") -> str:
    """
    清理原始檔名，只留下可讀的短字串。
    實際儲存檔名仍會加上時間與雜湊，避免直接使用使用者原始檔名。
    """
    stem = Path(str(filename or default_name)).stem.strip() or default_name
    stem = re.sub(r'[\\/:*?"<>|]+', "_", stem)
    stem = re.sub(r"\s+", "_", stem).strip("._- ")
    stem = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_\-]+", "_", stem)
    return (stem[:32] or default_name)


def safe_output_child_path(folder: Path, filename: str) -> Path:
    """
    防止路徑穿越：確認輸出位置一定在指定資料夾裡。
    """
    folder.mkdir(parents=True, exist_ok=True)
    target = (folder / filename).resolve()
    base = folder.resolve()
    if base != target.parent and base not in target.parents:
        raise ValueError("檔案路徑不安全，已阻擋寫入。")
    return target


def validate_uploaded_file(uploaded_file: Any, allowed_exts: set[str], max_bytes: int, label: str) -> tuple[str, bytes]:
    suffix = Path(str(getattr(uploaded_file, "name", ""))).suffix.lower()
    if suffix not in allowed_exts:
        raise ValueError(f"{label} 副檔名不允許：{suffix or '無副檔名'}")

    size = get_uploaded_file_size(uploaded_file)
    if size <= 0:
        raise ValueError(f"{label} 檔案大小異常。")
    if size > max_bytes:
        raise ValueError(f"{label} 檔案超過大小限制。")

    data = bytes(uploaded_file.getbuffer())
    return suffix, data


def validate_image_bytes(data: bytes) -> None:
    """
    圖片上傳除了副檔名，也用 Pillow 嘗試驗證內容。
    """
    try:
        with Image.open(BytesIO(data)) as img:
            img.verify()
    except Exception as e:
        raise ValueError("圖片內容無法驗證，可能不是有效圖片。") from e


def validate_doc_upload_bytes(data: bytes, suffix: str) -> None:
    """
    .docx 本質是 zip 結構，所以先做基本 zip 格式檢查。
    .doc 是舊 Word 格式，後續交給 Microsoft Word COM 轉檔。
    """
    if suffix == ".docx":
        try:
            if not zipfile.is_zipfile(BytesIO(data)):
                raise ValueError("DOCX 檔案格式驗證失敗。")
        except Exception as e:
            raise ValueError("DOCX 檔案格式驗證失敗。") from e


def build_secure_saved_filename(uploaded_name: str, suffix: str, data: bytes, prefix: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    digest = short_hash_bytes(data)
    stem = safe_upload_stem(uploaded_name, prefix)
    return f"{ts}_{digest}_{stem}{suffix}"


def save_uploaded_files(uploaded_files: list[Any]) -> list[Path]:
    """
    第九十八版：
    照片上傳安全化：
    - 只允許 IMAGE_EXTS
    - 限制大小
    - 用 Pillow 驗證圖片內容
    - 儲存檔名改成 時間 + 雜湊 + 清理後短檔名
    - 防止路徑穿越
    """
    upload_cache_dir().mkdir(parents=True, exist_ok=True)
    prune_upload_cache_to_current_df()
    paths: list[Path] = []
    errors: list[str] = []

    for f in uploaded_files:
        try:
            suffix, data = validate_uploaded_file(f, IMAGE_EXTS, MAX_IMAGE_UPLOAD_BYTES, "照片")
            validate_image_bytes(data)
            safe_name = build_secure_saved_filename(getattr(f, "name", "photo"), suffix, data, "photo")
            out = safe_output_child_path(upload_cache_dir(), safe_name)
            out.write_bytes(data)
            paths.append(out)
        except Exception as e:
            errors.append(f"{Path(str(getattr(f, 'name', '未命名'))).name}：{e}")

    if errors:
        st.error("部分照片未匯入：\n" + "\n".join(errors))

    return paths


def sanitize_upload_filename(filename: str, default_name: str = "抽查紀錄表樣板") -> str:
    """
    保留給舊函式相容使用。
    第八十七版主要改由 build_secure_saved_filename 產生安全檔名。
    """
    p = Path(str(filename or default_name))
    suffix = p.suffix.lower()
    if suffix not in allowed_record_doc_exts():
        suffix = ".docx"
    return f"{safe_upload_stem(filename, default_name)}{suffix}"


def unique_template_path(folder: Path, filename: str) -> Path:
    """
    不覆蓋既有檔案，避免 Word 開著同名檔時 Permission denied。
    """
    folder.mkdir(parents=True, exist_ok=True)
    safe_name = sanitize_upload_filename(filename)
    candidate = folder / safe_name

    if not candidate.exists():
        return candidate

    stem = candidate.stem
    suffix = candidate.suffix
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    for i in range(1, 1000):
        new_candidate = folder / f"{stem}_{ts}_{i}{suffix}"
        if not new_candidate.exists():
            return new_candidate

    return folder / f"{stem}_{ts}_{file_hash(candidate)[:8]}{suffix}"


def convert_doc_to_docx_with_libreoffice(doc_path: Path, docx_path: Path) -> Path:
    converter = find_libreoffice_binary()
    if not converter:
        raise RuntimeError(
            "偵測到 .doc 舊格式，但雲端環境尚未安裝 LibreOffice，無法自動轉成 .docx。"
            "請確認 GitHub 根目錄已加入 packages.txt，內容包含 libreoffice。"
        )

    work_dir = doc_path.parent / f"_doc_convert_{uuid.uuid4().hex}"
    out_dir = work_dir / "out"
    profile_dir = work_dir / "profile"
    out_dir.mkdir(parents=True, exist_ok=True)
    profile_dir.mkdir(parents=True, exist_ok=True)

    try:
        cmd = [
            converter,
            f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
            "--headless",
            "--nologo",
            "--nolockcheck",
            "--nodefault",
            "--nofirststartwizard",
            "--convert-to",
            "docx",
            "--outdir",
            str(out_dir),
            str(doc_path.resolve()),
        ]
        completed = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

        converted_path = out_dir / f"{doc_path.stem}.docx"
        if not converted_path.exists():
            candidates = sorted(out_dir.glob("*.docx"))
            converted_path = candidates[0] if candidates else converted_path

        if completed.returncode != 0 or not converted_path.exists():
            details = (completed.stderr or completed.stdout or "").strip()
            if details:
                details = f" LibreOffice 訊息：{details[:500]}"
            raise RuntimeError(f"LibreOffice 轉檔失敗。{details}")

        if docx_path.exists():
            docx_path = unique_template_path(docx_path.parent, docx_path.name)

        shutil.move(str(converted_path), str(docx_path))

        if not docx_path.exists():
            raise RuntimeError("LibreOffice 已執行，但沒有產生 .docx 檔案。")

        return docx_path

    except subprocess.TimeoutExpired as e:
        raise RuntimeError(".doc 轉 .docx 超過 120 秒，請確認檔案沒有損壞或過大。") from e

    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


def convert_doc_to_docx(doc_path: Path) -> Path:
    """
    第九十八版：
    支援 Word 97-2003 的 .doc 舊格式。
    修正 Streamlit 執行緒使用 Word COM 時出現：
    CoInitialize 尚未被呼叫 / 尚未呼叫 CoInitialize

    原因：
    Streamlit 不是在一般主執行緒呼叫 Word COM，
    所以每次轉檔前必須先 pythoncom.CoInitialize()。
    """
    base_docx = doc_path.with_suffix(".docx")
    if base_docx.exists():
        docx_path = unique_template_path(base_docx.parent, base_docx.name)
    else:
        docx_path = base_docx

    if os.name != "nt":
        return convert_doc_to_docx_with_libreoffice(doc_path, docx_path)

    try:
        import pythoncom  # type: ignore
        import win32com.client  # type: ignore
    except Exception as e:
        if find_libreoffice_binary():
            return convert_doc_to_docx_with_libreoffice(doc_path, docx_path)
        raise RuntimeError(
            "偵測到你上傳的是 .doc 舊格式，但目前環境沒有 pywin32，無法自動轉檔。"
            "請確認執行指令有加入 --with pywin32。"
        ) from e

    word = None
    document = None
    com_initialized = False

    try:
        pythoncom.CoInitialize()
        com_initialized = True

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        # FileFormat=16 是 wdFormatXMLDocument，也就是 .docx
        document = word.Documents.Open(str(doc_path.resolve()))
        document.SaveAs2(str(docx_path.resolve()), FileFormat=16)
        document.Close(False)
        document = None

        word.Quit()
        word = None

        if not docx_path.exists():
            raise RuntimeError("Word 轉檔完成但找不到輸出的 .docx 檔。")

        return docx_path

    except Exception as e:
        try:
            if document is not None:
                document.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass

        raise RuntimeError(
            f".doc 自動轉 .docx 失敗：{e}。"
            "請確認 Microsoft Word 已安裝、檔案沒有被 Word 開啟，"
            "如果 Word 有跳出保護檢視或安全性提示，請先關掉後重新拖入。"
        ) from e

    finally:
        if com_initialized:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

def clear_record_template_dir() -> None:
    """
    第九十八版：
    每次使用者重新匯入抽查紀錄表 Word，就清空舊樣本。
    這樣不會一直累積重複 Word，選單也不會越來越多、越來越慢。
    """
    record_template_dir().mkdir(parents=True, exist_ok=True)
    for p in record_template_dir().glob("*"):
        if p.is_file() and not p.name.startswith("~$") and p.suffix.lower() in [".doc", ".docx"]:
            try:
                p.unlink()
            except Exception:
                pass


def record_upload_batch_key(uploaded_files: list[Any]) -> str:
    parts = []
    for f in uploaded_files or []:
        parts.append(f"{Path(f.name).name}:{getattr(f, 'size', 0)}")
    return "|".join(parts)


def reset_record_template_uploader() -> None:
    """
    第九十八版：
    Streamlit 的多檔上傳在 accept_multiple_files=True 時，
    使用者再次按 Upload 可能會把新檔案加到舊清單後面。
    所以抽查紀錄表樣本處理完成後，立刻換 file_uploader key 讓它清空。

    目的：
    下一次拖入新的抽查紀錄表 Word，就是完整取代上一批，
    不會把上一個案子的樣本一起帶進來。
    """
    st.session_state["record_template_uploader_nonce"] = st.session_state.get("record_template_uploader_nonce", 0) + 1


def save_record_template_files(uploaded_files: list[Any], replace_mode: bool = False) -> list[Path]:
    """
    第九十八版：
    抽查紀錄表樣本依照「每次匯入的 Word」為準。

    修正：
    1. 重新匯入新 Word 時，先清掉舊樣本，只保留本次匯入。
    2. 處理完成後會重置 file_uploader，避免 Streamlit 多檔上傳把舊檔案留在清單裡。
    3. 選單只會依本次匯入 Word 的目錄工項產生。
    """
    record_template_dir().mkdir(parents=True, exist_ok=True)

    if not uploaded_files:
        return []

    batch_key = record_upload_batch_key(uploaded_files)
    if (not replace_mode) and st.session_state.get("record_template_upload_key") == batch_key:
        # 同一批檔案已經處理過，避免 rerun 時一直複製/轉檔。
        return []

    # 重要：
    # 每次匯入抽查紀錄表樣本，都先清掉舊樣本。
    # 這樣不會混到上一個案子的抽查表。
    clear_record_template_dir()

    saved: list[Path] = []
    errors: list[str] = []

    for f in uploaded_files:
        try:
            suffix, data = validate_uploaded_file(f, allowed_record_doc_exts(), MAX_DOC_UPLOAD_BYTES, "抽查紀錄表 Word")
            validate_doc_upload_bytes(data, suffix)

            safe_name = build_secure_saved_filename(getattr(f, "name", "record_template"), suffix, data, "record")
            raw_path = safe_output_child_path(record_template_dir(), safe_name)
            raw_path.write_bytes(data)

            if suffix == ".docx":
                saved.append(raw_path)
                continue

            # .doc 舊格式：轉成 .docx 後，刪掉原 .doc，避免資料夾出現兩份。
            converted = convert_doc_to_docx(raw_path)
            saved.append(converted)

            try:
                raw_path.unlink()
            except Exception:
                pass

        except Exception as e:
            errors.append(f"{Path(str(getattr(f, 'name', '未命名'))).name}：{e}")

    st.session_state["record_template_upload_key"] = batch_key
    st.session_state.pop("record_work_items_cache", None)

    # 立刻重建快取，後面每張照片 selectbox 不用一直重讀 Word。
    try:
        st.session_state["record_work_items_cache"] = build_dynamic_record_work_items()
    except Exception:
        st.session_state.pop("record_work_items_cache", None)

    if errors:
        if supports_legacy_doc_conversion():
            hint = "處理方式：確認該檔案沒有被 Word 開啟，並確認執行指令有 --with pywin32；關掉 Word 後重新拖入。"
        else:
            hint = "處理方式：請確認 GitHub 根目錄有 packages.txt 且內容包含 libreoffice；更新後在 Streamlit Cloud 重新啟動 app。"
        st.error(
            "有樣本匯入或轉檔失敗：\n"
            + "\n".join(errors)
            + f"\n\n{hint}"
        )

    return saved

def list_record_templates() -> list[Path]:
    if not record_template_dir().exists():
        return []
    # 後續套版只使用 .docx；.doc 上傳後會自動轉成 .docx。
    files = [
        p for p in record_template_dir().glob("*.docx")
        if not p.name.startswith("~$")
    ]
    files.sort(key=lambda p: p.name)
    return files


def parse_work_item_code(label: str) -> str:
    text = str(label or "").upper()
    m = re.search(r"B\d{2}", text)
    return m.group(0) if m else ""


def normalize_record_item_name(name: str) -> str:
    text = normalize_text(str(name or ""))
    text = re.sub(r"\s+", "", text)
    text = text.replace("抽查程序流程圖名稱", "")
    text = text.strip("：:｜|-_ ")
    return text


def is_record_catalog_table(table) -> bool:
    """
    判斷這張表是不是「工程抽查紀錄表目錄」。
    必須像目錄表，才拿來產生下拉工項。
    避免把後面正式 B05 / B06 抽查表本體也誤抓成工項，造成重複。
    """
    try:
        text = normalize_text("\n".join(cell.text for row in table.rows for cell in row.cells))
    except Exception:
        return False

    code_count = len(set(re.findall(r"\bB\d{2}\b", text, flags=re.IGNORECASE)))

    # 最準：使用者目錄表通常有這些欄位。
    if "抽查程序流程圖名稱" in text and ("序號" in text or "頁碼" in text):
        return True

    # 次準：有很多 Bxx + 頁碼/附錄，才當目錄。
    if code_count >= 3 and "抽查紀錄表" in text and ("頁碼" in text or "附錄" in text):
        return True

    return False


def extract_items_from_catalog_table(table, docx_path: Path) -> list[dict[str, Any]]:
    """
    只從目錄表列出工項。
    """
    items: list[dict[str, Any]] = []

    for row in table.rows:
        cells = [normalize_text(c.text) for c in row.cells]
        row_text = " ".join(cells)
        m = re.search(r"\b(B\d{2})\b", row_text, re.IGNORECASE)
        if not m:
            continue

        code = m.group(1).upper()
        name = ""
        page = ""

        for c in cells:
            if "抽查紀錄表" in c:
                name = c.strip()
                break

        if not name:
            candidates = []
            for c in cells:
                cc = c.strip()
                if not cc:
                    continue
                if re.fullmatch(r"B\d{2}", cc, re.IGNORECASE):
                    continue
                if "序號" in cc or "頁碼" in cc or "抽查程序流程圖名稱" in cc:
                    continue
                if "附錄" in cc:
                    page = cc
                    continue
                candidates.append(cc)
            if candidates:
                name = max(candidates, key=len)

        for c in cells:
            if "附錄" in c:
                page = c.strip()

        name = normalize_record_item_name(name) or f"{code}抽查紀錄表"
        label = f"{code} {name}"

        items.append({
            "code": code,
            "name": name,
            "page": page,
            "label": label,
            "template_path": str(docx_path),
            "source_file": docx_path.name,
        })

    return items


def dedupe_work_items_by_code(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """
    工項就是目錄那幾個。
    同一個 B05 / B07 如果被解析到多次，只保留第一筆。
    """
    output: list[dict[str, Any]] = []
    seen_codes: set[str] = set()
    seen_labels: set[str] = set()

    for item in items:
        code = str(item.get("code", "")).strip().upper()
        label = str(item.get("label", "")).strip()

        # 有 Bxx 時以 Bxx 去重；沒有 Bxx 時用 label 去重。
        if code:
            if code in seen_codes:
                continue
            seen_codes.add(code)
        else:
            if label in seen_labels:
                continue
            seen_labels.add(label)

        output.append(item)

    return output


def extract_work_items_from_docx(docx_path: Path) -> list[dict[str, Any]]:
    """
    第九十八版：
    抽查表工項只依照 Word 裡的「目錄表」產生。

    修正前的問題：
    程式掃到後面正式抽查表本體，例如：
    - B05 護岸工程抽查紀錄表
    - B05 護岸工程施工抽查紀錄表
    所以 UI 會出現兩個 B05。

    修正後：
    只讀目錄表，不讀後面正式表格本體。
    重新匯入 Word 時，仍會取代上一批樣本。
    """
    items: list[dict[str, Any]] = []

    try:
        from docx import Document
        doc = Document(str(docx_path))

        catalog_tables = [table for table in doc.tables if is_record_catalog_table(table)]

        if catalog_tables:
            for table in catalog_tables:
                items.extend(extract_items_from_catalog_table(table, docx_path))

            items = dedupe_work_items_by_code(items)
            return items

    except Exception:
        pass

    # 若 Word 真的沒有目錄表，才退回用檔名建立一個工項。
    stem = docx_path.stem
    code = parse_work_item_code(stem)
    name = normalize_record_item_name(stem)
    if not name:
        name = f"{code}抽查紀錄表" if code else stem
    label = f"{code} {name}".strip()

    return [{
        "code": code,
        "name": name,
        "page": "",
        "label": label,
        "template_path": str(docx_path),
        "source_file": docx_path.name,
    }]

def build_dynamic_record_work_items() -> list[dict[str, Any]]:
    """
    真正掃描 Word 的函式。
    注意：只在匯入後或快取不存在時執行，不會每張照片都重掃。
    """
    all_items: list[dict[str, Any]] = []
    seen_codes: set[str] = set()
    seen_labels: set[str] = set()

    for p in list_record_templates():
        for item in extract_work_items_from_docx(p):
            label = str(item.get("label", "")).strip()
            code = str(item.get("code", "")).strip().upper()
            if not label:
                continue

            # 工項就是目錄那幾個。
            # 同一個 B05 / B07 不管名稱略有不同，都只保留第一個。
            if code:
                if code in seen_codes:
                    continue
                seen_codes.add(code)
            else:
                if label in seen_labels:
                    continue
                seen_labels.add(label)

            all_items.append(item)

    def sort_key(item: dict[str, Any]):
        code = str(item.get("code", ""))
        m = re.search(r"(\d+)", code)
        num = int(m.group(1)) if m else 999
        return (num, str(item.get("label", "")))

    all_items.sort(key=sort_key)
    return all_items


def get_dynamic_record_work_items() -> list[dict[str, Any]]:
    """
    UI 下拉選單使用快取，避免每次重跑都讀 Word，提升速度。
    """
    cached = st.session_state.get("record_work_items_cache")
    if isinstance(cached, list):
        return cached

    items = build_dynamic_record_work_items()
    st.session_state["record_work_items_cache"] = items
    return items


def option_label_from_record_item(item: dict[str, Any]) -> str:
    # 選單只顯示目錄工項，不顯示檔名、不顯示重複來源。
    return str(item.get("label", "") or item.get("name", "")).strip()


def record_template_options() -> list[str]:
    items = get_dynamic_record_work_items()
    if not items:
        return ["請先匯入抽查紀錄表 Word"]

    return ["請選擇抽查表工項"] + [option_label_from_record_item(item) for item in items]


def get_record_item_by_label(label: str) -> dict[str, Any] | None:
    label = str(label or "").strip()
    for item in get_dynamic_record_work_items():
        if label == option_label_from_record_item(item):
            return item
    return None


def record_template_path_from_label(label: str) -> str:
    if str(label or "") in ["", "請選擇抽查表工項", "請先匯入抽查紀錄表 Word"]:
        return ""

    item = get_record_item_by_label(label)
    if item:
        return str(item.get("template_path", ""))

    return ""


def record_work_item_code_from_label(label: str) -> str:
    item = get_record_item_by_label(label)
    if item:
        return str(item.get("code", ""))
    return parse_work_item_code(label)


def record_work_item_name_from_label(label: str) -> str:
    item = get_record_item_by_label(label)
    if item:
        return str(item.get("name", ""))
    return str(label or "")


def missing_selected_record_templates(df: pd.DataFrame) -> list[str]:
    if df.empty or "抽查表工項" not in df.columns:
        return []

    missing: list[str] = []
    selected = df[df.get("輸出", True) == True].copy()

    for _, row in selected.iterrows():
        label = str(row.get("抽查表工項", "") or "").strip()
        if not label:
            continue

        path = str(row.get("抽查表工項路徑", "") or row.get("抽查表樣本路徑", "") or "").strip()
        if not path or not Path(path).exists():
            if label not in missing:
                missing.append(label)

    return missing


# 第九十八版：已移除 自動辨識功能。
# 使用者改為手動填寫日期、檢查項目、設計值、實測值。
# 保留後續表單輸出流程，不再呼叫 rapidocr_onnxruntime。


def save_user_memory(df: pd.DataFrame) -> None:
    """
    第九十八版：
    自動辨識記憶庫已移除。
    使用者手動填寫欄位，不新增額外輸入選單。
    """
    return


def normalize_rotation(value: Any) -> int:
    """
    照片旋轉角度只允許 0 / 90 / 180 / 270。
    """
    try:
        angle = int(value or 0)
    except Exception:
        angle = 0
    return angle % 360


@st.cache_data(show_spinner=False, max_entries=512)
def build_display_image_bytes(path_str: str, rotate_angle: int, mtime_ns: int, file_size: int) -> bytes:
    """
    第九十八版：
    UI 預覽用快取縮圖。
    之前每次 rerun 都重新讀取原始大圖，照片多時會 lag。
    這版改成：
    - 只在照片檔案改變或旋轉角度改變時重新產生縮圖
    - UI 顯示縮圖，不改原始照片
    """
    path = Path(path_str)
    img = safe_open_image(path)

    angle = normalize_rotation(rotate_angle)
    if angle:
        img = img.rotate(-angle, expand=True)

    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")

    # UI 預覽不需要原始大圖尺寸，縮到合理大小可大幅減少卡頓。
    img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)

    bio = BytesIO()
    img.save(bio, format="JPEG", quality=85, optimize=True)
    return bio.getvalue()


def build_display_image(path: Path, rotate_angle: int = 0) -> bytes:
    """
    UI 預覽用圖片。
    回傳快取後的 JPEG bytes，讓 st.image 顯示速度比較穩。
    """
    try:
        stat = path.stat()
        return build_display_image_bytes(str(path), normalize_rotation(rotate_angle), stat.st_mtime_ns, stat.st_size)
    except Exception:
        # 若快取縮圖失敗，保底回傳原本方式。
        img = safe_open_image(path)
        angle = normalize_rotation(rotate_angle)
        if angle:
            img = img.rotate(-angle, expand=True)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        bio = BytesIO()
        img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
        img.save(bio, format="JPEG", quality=85)
        return bio.getvalue()


def safe_output_name(name: str, default_name: str = "工程抽查照片_UI版") -> str:
    """
    第九十八版：
    使用者可自行輸入輸出 Word 檔案名稱。
    這裡會移除 Windows 不允許的檔名字元。
    """
    text = str(name or "").strip()
    if not text:
        text = default_name

    text = re.sub(r'[\\/:*?"<>|]+', "_", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.strip(" .")

    if not text:
        text = default_name

    if len(text) > 80:
        text = text[:80].rstrip()

    return text


def make_labeled_download_name(user_prefix: str, label: str, suffix: str = ".docx") -> str:
    prefix = safe_output_name(user_prefix, "").strip()
    label = safe_output_name(label, "Word").strip()

    if not prefix:
        stem = label
    elif prefix.endswith(label):
        stem = prefix
    else:
        stem = f"{prefix}{label}"

    return f"{stem}{suffix}"


def save_docx_to_bytes(doc) -> bytes:
    """
    把 Word 文件存到記憶體，供 st.download_button 下載。
    不在本機輸出資料夾留下 .docx 檔。
    """
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def _set_rfonts(rfonts, font_name: str = "標楷體") -> None:
    """
    Word 中文字型要同時設定 eastAsia。
    ascii/hAnsi/cs 也一起設定，避免 Word 顯示成新細明體或 MS Mincho。
    """
    for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
        try:
            rfonts.set(qn(f"w:{attr}"), font_name)
        except Exception:
            pass
    try:
        rfonts.set(qn("w:hint"), "eastAsia")
    except Exception:
        pass


def set_run_font_kai(run, font_name: str = "標楷體") -> None:
    """
    第九十八版：
    修正前版有些文字沒有真正套到標楷體的問題。
    原因是 qn 沒有在字型函式可用範圍內，導致 try/pass 靜默失敗。
    """
    try:
        run.font.name = font_name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        _set_rfonts(rFonts, font_name)
    except Exception:
        pass


def set_paragraph_font_kai(paragraph, font_name: str = "標楷體") -> None:
    for run in paragraph.runs:
        set_run_font_kai(run, font_name)


def set_table_font_kai(table, font_name: str = "標楷體") -> None:
    try:
        seen = set()
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for p in cell.paragraphs:
                    set_paragraph_font_kai(p, font_name)
                for nested_table in cell.tables:
                    set_table_font_kai(nested_table, font_name)
    except Exception:
        pass


def set_doc_default_font_kai(doc, font_name: str = "標楷體") -> None:
    """
    設定 Word 的文件預設字型。
    這會讓游標停在文件空白處或新增文字時，也顯示標楷體。
    """
    try:
        styles_element = doc.styles.element
        doc_defaults = styles_element.find(qn("w:docDefaults"))
        if doc_defaults is None:
            doc_defaults = OxmlElement("w:docDefaults")
            styles_element.insert(0, doc_defaults)

        rpr_default = doc_defaults.find(qn("w:rPrDefault"))
        if rpr_default is None:
            rpr_default = OxmlElement("w:rPrDefault")
            doc_defaults.append(rpr_default)

        rpr = rpr_default.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            rpr_default.append(rpr)

        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)

        _set_rfonts(rfonts, font_name)
    except Exception:
        pass


def set_all_style_fonts_kai(doc, font_name: str = "標楷體") -> None:
    """
    把文件內已存在的樣式都改成標楷體。
    """
    try:
        for style in doc.styles:
            try:
                style.font.name = font_name
                rPr = style._element.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                _set_rfonts(rFonts, font_name)
            except Exception:
                pass
    except Exception:
        pass


def force_all_existing_rfonts_kai(doc, font_name: str = "標楷體") -> None:
    """
    直接掃 Word XML 內所有 w:rFonts，全部改成標楷體。
    這是為了避免表格複製後仍沿用 MS Mincho / 新細明體。
    """
    try:
        for el in doc._element.iter():
            if el.tag == qn("w:rFonts"):
                _set_rfonts(el, font_name)
    except Exception:
        pass

    try:
        for part in [section.header for section in doc.sections] + [section.footer for section in doc.sections]:
            for el in part._element.iter():
                if el.tag == qn("w:rFonts"):
                    _set_rfonts(el, font_name)
    except Exception:
        pass


def set_document_font_kai(doc, font_name: str = "標楷體") -> None:
    """
    輸出 Word 全文件套用標楷體。
    第八十七版會同時處理：
    - 文件預設字型 docDefaults
    - 所有 styles
    - 所有 paragraph runs
    - 所有 table cell runs
    - header / footer
    - 既有 XML 內所有 w:rFonts
    """
    set_doc_default_font_kai(doc, font_name)
    set_all_style_fonts_kai(doc, font_name)

    for p in doc.paragraphs:
        set_paragraph_font_kai(p, font_name)

    for table in doc.tables:
        set_table_font_kai(table, font_name)

    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                set_paragraph_font_kai(p, font_name)
            for t in section.header.tables:
                set_table_font_kai(t, font_name)
            for p in section.footer.paragraphs:
                set_paragraph_font_kai(p, font_name)
            for t in section.footer.tables:
                set_table_font_kai(t, font_name)
    except Exception:
        pass

    force_all_existing_rfonts_kai(doc, font_name)


def find_first_docx(folder: Path) -> Path | None:
    if not folder.exists():
        return None
    files = [p for p in folder.glob("*.docx") if not p.name.startswith("~$")]
    files.sort(key=lambda p: p.name)
    return files[0] if files else None


def clear_cell(cell) -> None:
    cell.text = ""


def set_cell(cell, text: str, bold: bool = False) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(str(text or ""))
    run.bold = bold
    set_run_font_kai(run)


def append_table_from_template(doc, template_tbl_xml):
    from copy import deepcopy
    from docx.table import Table

    new_tbl_xml = deepcopy(template_tbl_xml)
    body = doc._body._element
    sect_pr = body.sectPr

    if sect_pr is not None:
        body.remove(sect_pr)
        body.append(new_tbl_xml)
        body.append(sect_pr)
    else:
        body.append(new_tbl_xml)

    return Table(new_tbl_xml, doc)


def prepare_image(path: Path, rotate_angle: int = 0) -> BytesIO:
    """
    第九十八版：
    照片轉成記憶體 BytesIO 給 Word 使用，不再輸出 _ui_word_images 暫存檔。

    修正重點：
    - 不再裁切照片。
    - 不再強制把照片 crop 成 1600 x 950。
    - 改成等比例縮小後放到白底畫布中央，確保整張照片完整保留。
    """
    img = safe_open_image(path)
    rotate_angle = normalize_rotation(rotate_angle)
    if rotate_angle:
        img = img.rotate(-rotate_angle, expand=True)

    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    if img.mode == "L":
        img = img.convert("RGB")

    target_w, target_h = 1600, 950

    # 等比例縮放到指定範圍內，不裁切、不變形。
    img.thumbnail((target_w, target_h), Image.Resampling.LANCZOS)

    # 用白底補滿固定比例，讓 Word 表格版面穩定。
    canvas = Image.new("RGB", (target_w, target_h), "white")
    left = (target_w - img.width) // 2
    top = (target_h - img.height) // 2
    canvas.paste(img, (left, top))

    bio = BytesIO()
    canvas.save(bio, "JPEG", quality=92)
    bio.seek(0)
    return bio

def blank_photo_block(table, block_index: int) -> None:
    """把某個照片區塊留成空白：保留日期／檢查項目／設計值／實測值的欄位標題文字，
    但不填任何數值，照片格也留空。方便日後在 Word 裡手動補資料。"""
    if block_index == 0:
        row_date, row_item, row_value, row_photo = 0, 1, 2, 3
        design_col, actual_col = 0, 2
    else:
        row_date, row_item, row_value, row_photo = 4, 5, 6, 7
        design_col, actual_col = 0, 1

    set_cell(table.cell(row_date, 0), "日期：")
    set_cell(table.cell(row_item, 0), "檢查項目：")
    set_cell(table.cell(row_value, design_col), "設計值：")
    set_cell(table.cell(row_value, actual_col), "實測值：")
    clear_cell(table.cell(row_photo, 0))


def fill_photo_block(table, block_index: int, row: pd.Series) -> int | None:
    """
    填完這個照片區塊後，回傳「設計值/實測值那一列」的列索引；
    若該列兩個值都沒填，呼叫端稍後會把這個索引整列從表格刪除。
    回傳 None 代表這一列有值，不需要刪除。
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    if block_index == 0:
        row_date, row_item, row_value, row_photo = 0, 1, 2, 3
        design_col, actual_col = 0, 2
    else:
        row_date, row_item, row_value, row_photo = 4, 5, 6, 7
        design_col, actual_col = 0, 1

    design_text = str(row.get("設計值", "") or "").strip()
    actual_text = str(row.get("實測值", "") or "").strip()

    set_cell(table.cell(row_date, 0), f"日期：{row.get('日期', '')}")
    set_cell(table.cell(row_item, 0), f"檢查項目：{row.get('檢查項目', '')}")
    # 沒有填設計值/實測值時（例如材料進場照片），Word 裡不顯示這兩個欄位標籤。
    set_cell(table.cell(row_value, design_col), f"設計值：{design_text}" if design_text else "")
    set_cell(table.cell(row_value, actual_col), f"實測值：{actual_text}" if actual_text else "")

    photo_cell = table.cell(row_photo, 0)
    clear_cell(photo_cell)
    p = photo_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
    except Exception:
        pass

    img_stream = prepare_image(Path(str(row.get("照片路徑", ""))), int(row.get("旋轉角度", 0) or 0))
    p.add_run().add_picture(img_stream, width=Inches(6.15))

    return row_value if not design_text and not actual_text else None


def remove_table_row_at(table, row_index: int) -> None:
    tbl = table._tbl
    rows = list(tbl.tr_lst)
    if 0 <= row_index < len(rows):
        tbl.remove(rows[row_index])


def set_template_header(doc, project: str, location: str) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()

        if text.startswith("工程名稱"):
            for run in p.runs:
                run.text = ""
            r = p.add_run(f"工程名稱：{project}")
            set_run_font_kai(r)
        elif text.startswith("施工地點") or text.startswith("分項工程"):
            label = "分項工程" if text.startswith("分項工程") else "施工地點"
            for run in p.runs:
                run.text = ""
            r = p.add_run(f"{label}：{location}")
            set_run_font_kai(r)


def add_repeated_photo_header(doc, project: str, location: str) -> None:
    """
    第九十八版：
    每一頁都補上：
    工程抽查照片
    工程名稱
    施工地點 / 分項工程
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("工程抽查照片")
    r.bold = True
    r.font.size = Pt(16)
    set_run_font_kai(r)

    p1 = doc.add_paragraph()
    r1 = p1.add_run(f"工程名稱：{project}")
    set_run_font_kai(r1)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(f"施工地點：{location}")
    set_run_font_kai(r2)


def create_photo_word(df: pd.DataFrame, output_file_name: str = '', base_photo_word_bytes: bytes | None = None) -> dict[str, Any]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    template = find_first_docx(PHOTO_TEMPLATE_DIR)
    selected = df[df["輸出"] == True].copy()

    if selected.empty:
        raise ValueError("沒有勾選任何要輸出的照片。")

    project = str(selected.iloc[0].get("工程名稱", ""))
    location = str(selected.iloc[0].get("施工地點", ""))

    append_to_existing = bool(base_photo_word_bytes)

    if append_to_existing:
        doc = Document(BytesIO(base_photo_word_bytes))

        if template and template.exists():
            template_doc = Document(str(template))
            if not template_doc.tables:
                raise RuntimeError("工程抽查照片樣板沒有表格，請放入正確的工程抽查表.docx。")
            template_tbl_xml = template_doc.tables[0]._tbl
        else:
            template_tbl_xml = None

    elif template and template.exists():
        doc = Document(str(template))
        set_template_header(doc, project, location)

        if not doc.tables:
            raise RuntimeError("工程抽查照片樣板沒有表格，請放入正確的工程抽查表.docx。")

        template_tbl_xml = doc.tables[0]._tbl

        for tbl in list(doc.tables):
            tbl._element.getparent().remove(tbl._element)

        remove_empty_body_paragraphs(doc)
        compact_photo_word_spacing(doc)
    else:
        doc = Document()
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = h.add_run("工程抽查照片")
        rh.bold = True
        rh.font.size = Pt(16)
        set_run_font_kai(rh)

        p_project = doc.add_paragraph(f"工程名稱：{project}")
        set_paragraph_font_kai(p_project)
        p_location = doc.add_paragraph(f"施工地點：{location}")
        set_paragraph_font_kai(p_location)
        template_tbl_xml = None
        compact_photo_word_spacing(doc)

    selected = sort_for_word_output(selected)
    rows = list(selected.iterrows())

    grouped_rows: list[list[tuple[Any, pd.Series]]] = []
    current_group: list[tuple[Any, pd.Series]] = []
    current_key = None

    for row_item in rows:
        _idx, row = row_item
        key = photo_word_page_group_key(row)
        if current_key is not None and key != current_key:
            grouped_rows.append(current_group)
            current_group = []
        current_group.append(row_item)
        current_key = key

    if current_group:
        grouped_rows.append(current_group)

    page_count = 0

    for group in grouped_rows:
        for start in range(0, len(group), 2):
            chunk = group[start:start + 2]

            if append_to_existing or page_count > 0:
                doc.add_page_break()
                add_repeated_photo_header(doc, project, location)

            _, row1 = chunk[0]

            if template_tbl_xml is not None:
                table = append_table_from_template(doc, template_tbl_xml)
                rows_to_remove = []

                empty_value_row = fill_photo_block(table, 0, row1)
                if empty_value_row is not None:
                    rows_to_remove.append(empty_value_row)

                if len(chunk) >= 2:
                    _, row2 = chunk[1]
                    empty_value_row = fill_photo_block(table, 1, row2)
                    if empty_value_row is not None:
                        rows_to_remove.append(empty_value_row)
                else:
                    # 只有一張照片時，下半格保留成完全空白的表格，
                    # 方便日後在 Word 裡手動補照片或資料。
                    blank_photo_block(table, 1)

                # 由列索引大到小刪，避免刪除前面的列導致後面列索引跑掉。
                for row_index in sorted(rows_to_remove, reverse=True):
                    remove_table_row_at(table, row_index)
            else:
                design_text = str(row1.get("設計值", "") or "").strip()
                actual_text = str(row1.get("實測值", "") or "").strip()
                doc.add_paragraph(f"日期：{row1.get('日期', '')}")
                doc.add_paragraph(f"檢查項目：{row1.get('檢查項目', '')}")
                value_parts = []
                if design_text:
                    value_parts.append(f"設計值：{design_text}")
                if actual_text:
                    value_parts.append(f"實測值：{actual_text}")
                if value_parts:
                    doc.add_paragraph("　".join(value_parts))
                doc.add_picture(prepare_image(Path(str(row1.get("照片路徑", ""))), int(row1.get("旋轉角度", 0) or 0)))

            page_count += 1

    file_name = make_labeled_download_name(output_file_name, "工程抽查照片")
    compact_photo_word_spacing(doc)
    set_document_font_kai(doc)
    force_photo_title_size_16(doc)

    return {
        "file_name": file_name,
        "data": save_docx_to_bytes(doc),
    }


def roc_text_to_date(value: str) -> date:
    """
    第九十八版：
    日期解析支援：
    - 115年5月19日
    - 2026/5/19
    - 2026-05-19
    - 114/12/1

    修正：
    舊版本可能把 2026/5/21 誤存成 126年5月21日。
    若偵測到民國年明顯落在未來 10 年以上，會自動往回修正 11 年：
    126年 → 115年。
    """
    text = normalize_text(str(value or ""))

    m = re.search(r"(?<!\d)(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", text)
    if not m:
        m = re.search(r"(?<!\d)(\d{4})\s*[/-]\s*(\d{1,2})\s*[/-]\s*(\d{1,2})(?!\d)", text)
    if m:
        try:
            y = int(m.group(1))
            mo = int(m.group(2))
            d = int(m.group(3))
            if y >= 1911:
                return date(y, mo, d)
        except Exception:
            pass

    m = re.search(r"(?<!\d)(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", text)
    if not m:
        m = re.search(r"(?<!\d)(\d{2,3})\s*[/-]\s*(\d{1,2})\s*[/-]\s*(\d{1,2})(?!\d)", text)

    if not m:
        return date.today()

    try:
        y = int(m.group(1))
        mo = int(m.group(2))
        d = int(m.group(3))

        if y < 100:
            y += 100

        current_roc = date.today().year - 1911

        # 舊 bug 修正：2026/5/21 會被吃成 026/5/21 → 126年。
        # 對目前工程照片來說，超過目前民國年 8 年以上通常就是誤判。
        if y > current_roc + 8 and (y - 11) >= 100:
            y -= 11

        return date(y + 1911, mo, d)
    except Exception:
        return date.today()


def date_to_roc_text(value: date) -> str:
    try:
        roc_year = int(value.year) - 1911
        return f"{roc_year}年{value.month}月{value.day}日"
    except Exception:
        return ""


def append_uploaded_photo_to_df(uploaded_file: Any, project: str, location: str, slot_key: str) -> None:
    """
    第九十八版：
    照片拖進表格後，只新增照片資料。
    不執行 自動辨識、不標記 自動辨識 狀態、不產生 自動辨識 欄位。
    """
    if uploaded_file is None:
        return

    uploaded_key = f"{slot_key}:{uploaded_file.name}:{getattr(uploaded_file, 'size', 0)}"
    state_key = f"slot_uploaded_key_{slot_key}"
    if st.session_state.get(state_key) == uploaded_key:
        return

    paths = save_uploaded_files([uploaded_file])
    if not paths:
        return

    # 只用使用者原始上傳檔名判斷日期；存檔檔名開頭是我們自己加的今天時間戳
    # （例如 20260701_...），不可以拿來當白板日期，否則會每張都變成今天。
    original_upload_name = str(getattr(uploaded_file, "name", "") or "")
    date_from_original_name = parse_date_from_text(original_upload_name)

    rows = []
    for path in paths:
        date_from_name = date_from_original_name

        rows.append({
            "輸出": True,
            "照片檔名": path.name,
            "照片路徑": str(path),
            "工程名稱": project or "",
            "施工地點": location or "",
            "施工廠商": st.session_state.get("project_contractor", ""),
            "日期": date_from_name,
            "照片分組": extract_group_name(path),
            "對應表單": "待確認",
            "抽查表工項": "",
            "抽查表工項路徑": "",
            "抽查表樣本路徑": "",
            "抽查表編號": "",
            "抽查表名稱": "",
            "旋轉角度": 0,
            "檢查項目": "",
            "設計值": "",
            "實測值": "",
            "圖片雜湊": file_hash(path),
        })

    new_df = pd.DataFrame(rows)

    if "df" in st.session_state and isinstance(st.session_state["df"], pd.DataFrame):
        st.session_state["df"] = pd.concat([st.session_state["df"], new_df], ignore_index=True)
    else:
        st.session_state["df"] = new_df

    st.session_state[state_key] = uploaded_key
    st.rerun()


def delete_one_photo(row_idx: int) -> None:
    """
    第九十八版：
    單一刪除照片。
    只從目前 UI / 輸出資料中移除，不會刪掉你電腦裡的原始照片檔。
    """
    if "df" not in st.session_state:
        return

    df = st.session_state["df"]
    if not isinstance(df, pd.DataFrame):
        return

    if row_idx not in df.index:
        return

    try:
        photo_path = Path(str(df.loc[row_idx].get("照片路徑", ""))).resolve()
        if photo_path.exists() and upload_cache_dir().resolve() in photo_path.parents:
            photo_path.unlink()
    except Exception:
        pass

    st.session_state["df"] = df.drop(index=row_idx).reset_index(drop=True)
    prune_upload_cache_to_current_df()
    st.rerun()


def render_photo_drop_zone(slot_key: str, project: str, location: str) -> None:
    """
    第九十八版：
    Upload 直接就是照片表格內的上傳區。
    不另外顯示黑框、不放在表格外。
    """
    uploaded_photo = st.file_uploader(
        "把照片拖到這格",
        type=["jpg", "jpeg", "png", "bmp", "tif", "tiff", "webp"],
        accept_multiple_files=False,
        key=f"slot_upload_{slot_key}",
    )
    append_uploaded_photo_to_df(uploaded_photo, project, location, slot_key)


def render_empty_photo_slot(slot_key: str, project: str = '', location: str = '') -> None:
    """
    第九十八版：
    空白照片格改成輕量版。
    以前每個空白格都建立 disabled 的 selectbox、text_area、text_input；
    照片格越多，按「新增照片表格」就越慢。
    這版空白格只保留上傳區，照片拖入後才顯示完整欄位。
    """
    with st.container(border=True):
        render_photo_drop_zone(slot_key, project, location)
        st.caption("照片拖入後，日期、抽查表工項、檢查項目、設計值、實測值才會顯示。")


def get_extra_empty_slots() -> int:
    try:
        return int(st.session_state.get("extra_empty_slots", 0))
    except Exception:
        return 0


def add_one_empty_slot() -> None:
    st.session_state["extra_empty_slots"] = get_extra_empty_slots() + 1


def fragment_if_available(func):
    """
    第九十八版：
    如果目前 Streamlit 版本支援 st.fragment，就讓照片區局部重跑。
    這樣按「新增照片表格」時，不必整個 App 重跑，比較不會卡。
    舊版 Streamlit 沒有 st.fragment 時，會自動維持原本行為。
    """
    fragment = getattr(st, "fragment", None)
    if callable(fragment):
        return fragment(func)
    return func


def render_dynamic_empty_slots(total_empty_slots: int, project: str = '', location: str = '') -> None:
    if total_empty_slots <= 0:
        return

    photos_per_page = 6
    for page_start in range(0, total_empty_slots, photos_per_page):
        batch_end = min(page_start + photos_per_page, total_empty_slots)

        top_cols = st.columns([1, 1, 1], gap="small")
        for i in range(3):
            slot_no = page_start + i
            with top_cols[i]:
                if slot_no < batch_end:
                    render_empty_photo_slot(f"empty_top_{page_start}_{i}", project, location)
                else:
                    st.empty()

        bottom_cols = st.columns([1, 1, 1], gap="small")
        for i in range(3):
            slot_no = page_start + 3 + i
            with bottom_cols[i]:
                if slot_no < batch_end:
                    render_empty_photo_slot(f"empty_bottom_{page_start}_{i}", project, location)
                else:
                    st.empty()

        if batch_end < total_empty_slots:
            st.divider()


def render_empty_grid_ui(project: str = '', location: str = '') -> None:
    st.markdown(
        """
        <style>
        .block-container { padding-top: 0.5rem; padding-left: 1rem; padding-right: 1rem; }
        div[data-testid="column"] { min-width: 0; }
        .photo-box { min-height: 180px; border: 2px dashed #cbd5e1; background: #f8fafc; margin-bottom: 0.35rem; }
        .date-title { font-weight: 900; font-size: 1.35rem; text-align: center; padding: 1.5rem 0 1rem; }
        .weekday { text-align: center; font-weight: 800; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    add_col1, add_col2 = st.columns([1, 2])
    with add_col1:
        st.button("新增照片表格", key="add_empty_slot_only", on_click=add_one_empty_slot, use_container_width=True)
    with add_col2:
        st.caption("照片超過 6 張時，可一直按『新增照片表格』增加空白格。")

    total_empty_slots = max(6, 6 + get_extra_empty_slots())
    render_dynamic_empty_slots(total_empty_slots, project, location)


render_empty_grid_ui = fragment_if_available(render_empty_grid_ui)


def render_photo_slot(edited: pd.DataFrame, row_idx: int, slot_key: str) -> pd.DataFrame:
    row = edited.loc[row_idx]
    img_path = Path(str(row.get("照片路徑", "")))
    key_base = f"{slot_key}_{row_idx}_{str(row.get('圖片雜湊', ''))[:8]}"

    with st.container(border=True):
        st.markdown(f"**照片 #{row_idx + 1}**")
        rotate_angle = normalize_rotation(row.get("旋轉角度", 0))

        if img_path.exists():
            st.image(build_display_image(img_path, rotate_angle), use_container_width=True)
        else:
            # 找不到照片時，也直接在照片格內提供上傳
            render_photo_drop_zone(f"{slot_key}_missing", "", "")

        rotate_col1, rotate_col2, rotate_col3 = st.columns([1, 1, 1], gap="small")
        with rotate_col1:
            if st.button("左轉90°", key=f"{key_base}_rotate_left", use_container_width=True):
                edited.loc[row_idx, "旋轉角度"] = normalize_rotation(rotate_angle - 90)
                st.session_state["df"] = edited
                st.rerun()
        with rotate_col2:
            if st.button("右轉90°", key=f"{key_base}_rotate_right", use_container_width=True):
                edited.loc[row_idx, "旋轉角度"] = normalize_rotation(rotate_angle + 90)
                st.session_state["df"] = edited
                st.rerun()
        with rotate_col3:
            if st.button("刪除", key=f"{key_base}_delete_photo", use_container_width=True):
                delete_one_photo(row_idx)

        template_options = record_template_options()
        default_template_option = template_options[0] if template_options else "請先匯入抽查紀錄表 Word"
        current_template = str(row.get("抽查表工項", "") or row.get("對應表單", "") or default_template_option)
        if current_template not in template_options:
            current_template = default_template_option

        selected_template = st.selectbox(
            "抽查表工項",
            options=template_options,
            index=template_options.index(current_template),
            key=f"{key_base}_record_template",
        )

        selected_template_path = record_template_path_from_label(selected_template)
        if selected_template not in ["請選擇抽查表工項", "請先匯入抽查紀錄表 Word"]:
            if selected_template_path:
                st.caption(f"已對應這次匯入的 Word：{Path(selected_template_path).name}")
            else:
                st.warning("尚未找到此工項的 Word 樣本，請先在左側匯入本案抽查紀錄表 Word。")

        current_date = str(row.get("日期", "") or "")
        default_pick = roc_text_to_date(current_date) if current_date else None

        # 第九十八版：
        # 移除額外確認按鈕。
        # 直接使用 Streamlit 原生日期欄位：點日期欄位會開月曆，選完日期就立即套用。
        # 第一百零一版：日期欄位不再預設今天，避免使用者忘記手動選白板日期時
        # 誤用今天日期；欄位留空會強制使用者自己選擇。
        picked_date = st.date_input(
            "日期",
            value=default_pick,
            key=f"{key_base}_date_picker_direct",
        )
        date_value = date_to_roc_text(picked_date) if picked_date else ""

        # 第一百版：畫面只保留使用者原本要填的欄位，不再新增額外控制項。
        item_value = st.text_area(
            "檢查項目",
            value=str(row.get("檢查項目", "") or ""),
            height=58,
            key=f"{key_base}_item",
        )

        c_design, c_actual = st.columns(2, gap="small")
        with c_design:
            design_value = st.text_input(
                "設計值",
                value=str(row.get("設計值", "") or ""),
                key=f"{key_base}_design",
            )
        with c_actual:
            actual_value = st.text_input(
                "實測值",
                value=str(row.get("實測值", "") or ""),
                key=f"{key_base}_actual",
            )

        if selected_template in ["請選擇抽查表工項", "請先匯入抽查紀錄表 Word"]:
            edited.loc[row_idx, "抽查表工項"] = ""
        else:
            edited.loc[row_idx, "抽查表工項"] = selected_template

        edited.loc[row_idx, "抽查表工項路徑"] = selected_template_path
        edited.loc[row_idx, "抽查表樣本路徑"] = selected_template_path
        edited.loc[row_idx, "抽查表編號"] = record_work_item_code_from_label(selected_template)
        edited.loc[row_idx, "抽查表名稱"] = record_work_item_name_from_label(selected_template)
        edited.loc[row_idx, "對應表單"] = edited.loc[row_idx, "抽查表工項"] or "待確認"
        edited.loc[row_idx, "旋轉角度"] = rotate_angle
        edited.loc[row_idx, "日期"] = date_value
        edited.loc[row_idx, "檢查項目"] = item_value
        edited.loc[row_idx, "設計值"] = design_value
        edited.loc[row_idx, "實測值"] = actual_value
        edited.loc[row_idx, "輸出"] = True

    return edited


def render_photo_card_editor(df: pd.DataFrame, project: str = '', location: str = '') -> pd.DataFrame:
    edited = df.copy().reset_index(drop=True)

    st.markdown(
        """
        <style>
        .block-container { padding-top: 0.5rem; padding-left: 1rem; padding-right: 1rem; }
        div[data-testid="column"] { min-width: 0; }
        .photo-box { min-height: 180px; border: 2px dashed #cbd5e1; background: #f8fafc; margin-bottom: 0.35rem; }
        .date-title { font-weight: 900; font-size: 1.35rem; text-align: center; padding: 1.5rem 0 1rem; }
        .weekday { text-align: center; font-weight: 800; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    if "輸出" not in edited.columns:
        edited["輸出"] = True
    if "對應表單" not in edited.columns:
        edited["對應表單"] = "待確認"
    if "抽查表工項" not in edited.columns:
        edited["抽查表工項"] = ""
    if "抽查表工項路徑" not in edited.columns:
        edited["抽查表工項路徑"] = ""
    if "抽查表樣本路徑" not in edited.columns:
        edited["抽查表樣本路徑"] = ""
    if "抽查表編號" not in edited.columns:
        edited["抽查表編號"] = ""
    if "抽查表名稱" not in edited.columns:
        edited["抽查表名稱"] = ""
    if "施工廠商" not in edited.columns:
        edited["施工廠商"] = st.session_state.get("project_contractor", "")
    if "旋轉角度" not in edited.columns:
        edited["旋轉角度"] = 0

    add_col1, add_col2 = st.columns([1, 2])
    with add_col1:
        st.button("新增照片表格", key="add_photo_slot_with_data", on_click=add_one_empty_slot, use_container_width=True)
    with add_col2:
        st.caption("如果照片表格不夠用，可一直按『新增照片表格』，系統會持續增加空白格。")

    photos_per_page = 6
    existing_count = len(edited)
    extra_slots = get_extra_empty_slots()
    total_slots = max(6, existing_count + extra_slots)

    for page_start in range(0, total_slots, photos_per_page):
        batch_end = min(page_start + photos_per_page, total_slots)

        top_cols = st.columns([1, 1, 1], gap="small")
        for i in range(3):
            slot_no = page_start + i
            with top_cols[i]:
                if slot_no < batch_end:
                    if slot_no < existing_count:
                        edited = render_photo_slot(edited, slot_no, f"top_{page_start}_{i}")
                    else:
                        render_empty_photo_slot(f"blank_top_{page_start}_{i}", project, location)
                else:
                    st.empty()

        bottom_cols = st.columns([1, 1, 1], gap="small")
        for i in range(3):
            slot_no = page_start + 3 + i
            with bottom_cols[i]:
                if slot_no < batch_end:
                    if slot_no < existing_count:
                        edited = render_photo_slot(edited, slot_no, f"bottom_{page_start}_{i}")
                    else:
                        render_empty_photo_slot(f"blank_bottom_{page_start}_{i}", project, location)
                else:
                    st.empty()

        if batch_end < total_slots:
            st.divider()

    # fragment 局部重跑時，main 不一定會整段重跑；
    # 所以這裡直接把最新照片資料寫回 session_state。
    st.session_state["df"] = edited
    return edited


render_photo_card_editor = fragment_if_available(render_photo_card_editor)


@st.cache_data(show_spinner=False, max_entries=64)
def read_output_file_bytes(path_str: str, mtime_ns: int, file_size: int) -> bytes:
    """
    下載按鈕讀檔快取。
    避免每次畫面重跑都重新讀 Word 檔。
    """
    return Path(path_str).read_bytes()


def set_last_generated_words(photo_word: dict[str, Any], record_words: list[dict[str, Any]], missing_record_templates: list[str]) -> None:
    """
    第九十八版：
    產出後只把 Word bytes 存在 session_state 供下載。
    不再把 Word 寫進「輸出」資料夾。
    """
    st.session_state["last_generated_words"] = {
        "photo_word": photo_word or {},
        "record_words": list(record_words or []),
        "missing_record_templates": list(missing_record_templates or []),
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def render_last_generated_word_downloads() -> None:
    """
    固定顯示上一次產出的 Word 下載按鈕。
    第九十八版：下載來源為記憶體 bytes，不需要本機存檔。
    """
    info = st.session_state.get("last_generated_words")
    if not isinstance(info, dict):
        return

    photo_word = info.get("photo_word", {}) or {}
    record_words = info.get("record_words", []) or []
    missing = info.get("missing_record_templates", []) or []
    created_at = str(info.get("created_at", "") or "")

    has_photo = bool(photo_word.get("data")) and bool(photo_word.get("file_name"))
    valid_records = [w for w in record_words if w.get("data") and w.get("file_name")]

    if not has_photo and not valid_records:
        return

    st.success(f"已完成 Word 輸出。{created_at}")

    if has_photo:
        st.download_button(
            "下載 工程抽查照片 Word",
            data=photo_word["data"],
            file_name=photo_word["file_name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_photo_word_{photo_word['file_name']}",
            on_click="ignore",
            use_container_width=True,
        )

    if valid_records:
        st.write("已選工項抽查紀錄表 Word：")
        for w in valid_records:
            st.download_button(
                f"下載 {w['file_name']}",
                data=w["data"],
                file_name=w["file_name"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_record_word_{w['file_name']}",
                on_click="ignore",
                use_container_width=True,
            )

    if missing:
        st.warning(
            "以下工項已選擇，但找不到對應抽查紀錄表 Word 樣本，所以未產出該工項抽查紀錄表："
            + "、".join(missing)
        )
def validate_existing_docx_upload(uploaded_file: Any, label: str) -> bytes:
    """
    第九十八版：
    接續舊 Word 只支援 .docx。
    因為這裡是直接把 Word 載入記憶體當底稿，不再額外存成暫存檔，避免資料夾累積。
    """
    suffix, data = validate_uploaded_file(uploaded_file, {".docx"}, MAX_DOC_UPLOAD_BYTES, label)
    validate_doc_upload_bytes(data, suffix)
    return data


def render_existing_word_uploader(
    title: str,
    uploader_label: str,
    state_bytes_key: str,
    state_name_key: str,
    uploader_key: str,
) -> None:
    st.subheader(title)
    uploaded = st.file_uploader(
        uploader_label,
        type=["docx"],
        accept_multiple_files=False,
        key=uploader_key,
    )

    if uploaded is not None:
        try:
            data = validate_existing_docx_upload(uploaded, title)
            st.session_state[state_bytes_key] = data
            st.session_state[state_name_key] = Path(str(uploaded.name)).name
            st.success(f"已載入：{st.session_state[state_name_key]}")
        except Exception as e:
            st.session_state.pop(state_bytes_key, None)
            st.session_state.pop(state_name_key, None)
            st.error(f"{title} 載入失敗：{e}")

    elif st.session_state.get(state_bytes_key):
        st.success(f"目前接續：{st.session_state.get(state_name_key, '已匯入 Word')}")

    if st.session_state.get(state_bytes_key):
        if st.button(f"移除{title}", key=f"clear_{state_bytes_key}", use_container_width=True):
            st.session_state.pop(state_bytes_key, None)
            st.session_state.pop(state_name_key, None)
            st.rerun()


def get_shared_app_password() -> str:
    """
    讀取公司共用密碼。
    雲端部署請在 Streamlit Secrets 設定 APP_PASSWORD；
    本機測試也可用環境變數 APP_PASSWORD。
    """
    for key in ["APP_PASSWORD", "app_password", "shared_password"]:
        try:
            value = st.secrets.get(key, "")
            if str(value or "").strip():
                return str(value).strip()
        except Exception:
            pass

    for key in ["APP_PASSWORD", "STREAMLIT_APP_PASSWORD"]:
        value = os.environ.get(key, "")
        if str(value or "").strip():
            return str(value).strip()

    return ""


def require_shared_password_login() -> None:
    """
    第一版網站登入：一組公司共用密碼。
    密碼不寫死在程式碼，避免上 GitHub 或部署時外洩。
    """
    password = get_shared_app_password()

    if not password:
        st.error("尚未設定公司共用密碼，為了避免網站公開後外部人士直接使用，請先設定 APP_PASSWORD。")
        st.info("本機請建立 .streamlit/secrets.toml；雲端部署請在 Streamlit Cloud 的 Secrets 加上 APP_PASSWORD。")
        st.code('APP_PASSWORD = "請換成你的公司共用密碼"', language="toml")
        st.stop()

    if st.session_state.get("shared_password_authenticated") is True:
        with st.sidebar:
            st.caption("已通過公司共用密碼")
            if st.button("登出", key="logout_shared_password", use_container_width=True):
                st.session_state.pop("shared_password_authenticated", None)
                st.rerun()
            st.divider()
        return

    st.title("工程抽查 UI 工具")
    st.caption("請輸入公司共用密碼後使用。")

    with st.form("shared_password_login_form"):
        entered_password = st.text_input("共用密碼", type="password")
        submitted = st.form_submit_button("登入", use_container_width=True)

    if submitted:
        if hmac.compare_digest(str(entered_password or ""), password):
            st.session_state["shared_password_authenticated"] = True
            st.rerun()
        else:
            st.error("密碼錯誤，請再確認一次。")

    st.stop()


def main() -> None:
    st.set_page_config(
        page_title="工程抽查 UI 工具",
        page_icon="📋",
        layout="wide",
    )

    require_shared_password_login()
    ensure_session_runtime_paths()
    init_dirs()

    if "project_name" not in st.session_state:
        st.session_state["project_name"] = ""
    if "project_location" not in st.session_state:
        st.session_state["project_location"] = ""
    if "project_contractor" not in st.session_state:
        st.session_state["project_contractor"] = ""
    if "output_file_name" not in st.session_state:
        st.session_state["output_file_name"] = ""

    with st.sidebar:
        st.header("工程資料")
        st.caption("目前版本：第100版（無月曆套用按鈕／無最近下拉選單）")
        project = st.text_input(
            "工程名稱",
            placeholder="請輸入本案工程名稱，例如：哈拉灣溪橋下游右岸改善二期工程",
            key="project_name",
        )
        location = st.text_input(
            "施工地點 / 分項工程",
            placeholder="請輸入本案位置，例如：花蓮縣玉里鎮",
            key="project_location",
        )
        contractor = st.text_input(
            "施工廠商",
            placeholder="請輸入施工廠商名稱",
            key="project_contractor",
        )
        st.divider()
        output_file_name = st.text_input(
            "檔案名稱",
            placeholder="",
            help="可留空；空白時會自動使用工程抽查照片與抽查紀錄表作為檔名。",
            key="output_file_name",
        )

        clear_clicked = st.button(
            "開始新案件／清空暫存與輸出",
            help="只會清空目前使用者這個案件的暫存、輸出與匯入樣板。",
            use_container_width=True,
        )

        st.divider()
        st.header("接續既有 Word")
        st.caption("同一個案子隔幾天繼續貼照片時使用；換新案件時不要匯入舊 Word。")
        render_existing_word_uploader(
            "既有工程抽查照片 Word",
            "匯入舊照片 Word（.docx，可選）",
            "existing_photo_word_bytes",
            "existing_photo_word_name",
            "existing_photo_word_upload",
        )
        render_existing_word_uploader(
            "既有抽查紀錄表 Word",
            "匯入舊抽查表 Word（.docx，可選）",
            "existing_record_word_bytes",
            "existing_record_word_name",
            "existing_record_word_upload",
        )

        st.divider()
        st.header("工程抽查表樣板")
        st.caption("固定空白樣本。請先把固定樣板放在工程抽查照片樣板資料夾內。")
        current_photo_template = find_first_docx(PHOTO_TEMPLATE_DIR)
        if current_photo_template:
            st.success(f"目前樣板：{current_photo_template.name}")
        else:
            st.warning("尚未找到工程抽查表樣板。")

        st.divider()
        st.header("抽查紀錄表樣本")
        if supports_legacy_doc_conversion():
            st.caption("依不同案子匯入 Word。工項選單只會顯示這次匯入 Word 目錄裡的工項；重新匯入會取代上一批。")
        else:
            st.caption("雲端版支援 .doc/.docx；.doc 會先轉成 .docx。工項選單只會顯示這次匯入 Word 目錄裡的工項；重新匯入會取代目前使用者的上一批。")
        if "record_template_uploader_nonce" not in st.session_state:
            st.session_state["record_template_uploader_nonce"] = 0

        record_template_upload_key = f"record_template_uploads_{st.session_state['record_template_uploader_nonce']}"
        record_template_uploads = st.file_uploader(
            f"匯入抽查紀錄表樣本（{', '.join('.' + x for x in allowed_record_upload_types())}，可多選）",
            type=allowed_record_upload_types(),
            accept_multiple_files=True,
            key=record_template_upload_key,
        )
        if record_template_uploads:
            saved_templates = save_record_template_files(record_template_uploads, replace_mode=True)
            st.session_state.pop("last_generated_words", None)
            st.session_state["record_template_import_message"] = (
                f"已取代舊抽查紀錄表樣本，並只解析本次 Word 目錄表，目前只保留這次匯入的 {len(saved_templates)} 個 Word。"
                if saved_templates
                else "已清除舊抽查紀錄表樣本，但這次沒有成功匯入新的 Word。"
            )
            reset_record_template_uploader()
            st.rerun()

        import_message = st.session_state.pop("record_template_import_message", "")
        if import_message:
            st.success(import_message)

        record_templates = list_record_templates()
        if record_templates:
            dynamic_items = get_dynamic_record_work_items()
            with st.expander("目前匯入 Word 解析到的抽查表工項", expanded=False):
                if dynamic_items:
                    for item in dynamic_items:
                        st.write(f"- {option_label_from_record_item(item)}")
                else:
                    st.write("尚未解析到目錄工項。")
        else:
            st.info("尚未匯入抽查紀錄表 Word。")

    if clear_clicked:
        clear_current_case_files()
        reset_current_case_state()
        st.success("已清空本案暫存照片、輸出 Word、_ui_word_images，可開始新案件。")
        st.rerun()

    if "extra_empty_slots" not in st.session_state:
        st.session_state["extra_empty_slots"] = 0

    if "df" not in st.session_state:
        render_empty_grid_ui(project, location)
        return

    st.divider()
    df = st.session_state["df"].copy()
    edited = render_photo_card_editor(df, project, location)
    st.session_state["df"] = edited

    st.divider()
    st.subheader("③ 產出 Word")

    selected_count = int(edited["輸出"].fillna(False).sum())
    st.write(f"目前勾選輸出照片：**{selected_count}** 張")

    selected_mask = edited["輸出"].fillna(False) == True
    missing_date_rows = edited[selected_mask & (edited["日期"].fillna("").astype(str).str.strip() == "")]
    has_missing_date = len(missing_date_rows) > 0
    missing_photo_numbers = [idx + 1 for idx in missing_date_rows.index.tolist()]

    missing_item_rows = edited[selected_mask & (edited["抽查表工項"].fillna("").astype(str).str.strip() == "")]
    has_missing_item = len(missing_item_rows) > 0
    missing_item_numbers = [idx + 1 for idx in missing_item_rows.index.tolist()]

    c1, c2, c3 = st.columns([1, 1, 1])

    with c1:
        st.empty()

    with c2:
        make_words = st.button("產出 Word", type="primary", use_container_width=True)

    with c3:
        st.caption("本版只輸出 Word，不產生 Excel / ZIP。")

    if make_words:
        if selected_count == 0:
            st.error("至少要勾選一張照片才能產出 Word。")
        elif has_missing_date:
            numbers_text = "、".join(f"照片 #{n}" for n in missing_photo_numbers)
            st.error(
                f"無法產出：以下 {len(missing_photo_numbers)} 張照片還沒選日期，"
                f"請回到上面對應編號的照片把日期填好後再按產出 Word。\n\n未填日期：{numbers_text}"
            )
        elif has_missing_item:
            numbers_text = "、".join(f"照片 #{n}" for n in missing_item_numbers)
            st.error(
                f"無法產出：以下 {len(missing_item_numbers)} 張照片還沒選抽查表工項，"
                f"請回到上面對應編號的照片把抽查表工項選好後再按產出 Word。\n\n未選抽查表工項：{numbers_text}"
            )
        else:
            project_value = str(st.session_state.get("project_name", "") or project or "").strip()
            location_value = str(st.session_state.get("project_location", "") or location or "").strip()
            contractor_value = str(st.session_state.get("project_contractor", "") or contractor or "").strip()
            output_name_value = str(st.session_state.get("output_file_name", "") or output_file_name or "").strip()

            # 若使用者已經填在左側，但 Streamlit 當次 rerun 尚未同步，改用目前 DataFrame 裡的資料當 fallback。
            if not project_value and "工程名稱" in edited.columns and not edited.empty:
                project_value = str(edited["工程名稱"].dropna().astype(str).replace("", pd.NA).dropna().head(1).iloc[0]) if edited["工程名稱"].dropna().astype(str).str.strip().any() else ""
            if not location_value and "施工地點" in edited.columns and not edited.empty:
                location_value = str(edited["施工地點"].dropna().astype(str).replace("", pd.NA).dropna().head(1).iloc[0]) if edited["施工地點"].dropna().astype(str).str.strip().any() else ""
            if not contractor_value and "施工廠商" in edited.columns and not edited.empty:
                contractor_value = str(edited["施工廠商"].dropna().astype(str).replace("", pd.NA).dropna().head(1).iloc[0]) if edited["施工廠商"].dropna().astype(str).str.strip().any() else ""

            if not project_value or not location_value:
                st.error("請先在左側輸入工程名稱與施工地點 / 分項工程，再產出 Word。")
            else:
                try:
                    # 產出前強制以左側目前輸入的工程資料覆蓋勾選資料，避免不同案件沿用舊名稱。
                    edited.loc[edited["輸出"] == True, "工程名稱"] = project_value
                    edited.loc[edited["輸出"] == True, "施工地點"] = location_value
                    edited.loc[edited["輸出"] == True, "施工廠商"] = contractor_value

                    # 第九十八版：Word 只存在記憶體給使用者下載，不寫入輸出資料夾。
                    save_user_memory(edited)
                    photo_word = create_photo_word(
                        edited,
                        output_name_value,
                        base_photo_word_bytes=st.session_state.get("existing_photo_word_bytes"),
                    )
                    record_words = create_record_words(
                        edited,
                        output_name_value,
                        base_record_word_bytes=st.session_state.get("existing_record_word_bytes"),
                    )
                    missing_record_templates = missing_selected_record_templates(edited)

                    set_last_generated_words(photo_word, record_words, missing_record_templates)

                except Exception as e:
                    st.error(f"產出失敗：{e}")

    # 第九十八版：
    # 上一次產出的 Word 下載按鈕固定保留。
    # 按任一下載不會讓另一個下載按鈕消失。
    render_last_generated_word_downloads()


if __name__ == "__main__":
    main()
